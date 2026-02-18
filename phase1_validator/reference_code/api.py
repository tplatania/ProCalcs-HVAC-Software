"""
ProCalcs Dashboard API
Flask backend to serve dashboard and project analysis
"""

from flask import Flask, jsonify, request, redirect, session, send_from_directory, make_response
from flask_cors import CORS
import os
import json
import requests
from datetime import datetime, timedelta
from pathlib import Path
import anthropic
from google.cloud import firestore

# Initialize Firestore client
db = firestore.Client()

# Handle imports for both direct run and package run
try:
    from .config import Config
    from .project_analyzer import ProjectAnalyzer
    from .response_parser import parse_ai_analysis, generate_rfi_email
    from .outlook_integration import (
        get_auth_url, get_token_from_code, get_user_info, 
        send_email, get_rfi_email_template, get_followup_email_template, 
        get_review_request_template, read_emails, get_email_details,
        refresh_access_token as outlook_refresh_token
    )
    from .reference_api import reference_bp, init_reference_data
    from .quickbooks_routes import quickbooks_bp
    from .commission_routes import commission_bp
    from .zoom_routes import zoom_bp
    from .lead_events_routes import lead_events_bp
    from .self_check_routes import self_check_bp
    from .qc_routes import qc_bp
    from .zoho_oauth_routes import zoho_oauth_bp
    from .workdrive_routes import workdrive_bp
    from .email_routes import email_bp
    from .email_sending_routes import email_sending_bp
    from .zoho_routes import zoho_bp
    from .team_tasks_routes import team_tasks_bp
    from .user_signatures_routes import user_signatures_bp
    from .pipeline_routes import pipeline_bp, get_or_create_inquiry
    from . import zoho_integration as zoho
    from . import lead_events
    from . import event_detection
    from .quickbooks_service import get_qb_service
    from . import email_scanner
    from . import token_storage
    from . import user_roles
    from . import self_check_storage
    from . import official_qc_storage
    from . import google_sheets_service
    from . import streaming_analyzer
    from . import gemini_estimate
except ImportError:
    from config import Config
    from project_analyzer import ProjectAnalyzer
    from response_parser import parse_ai_analysis, generate_rfi_email
    from outlook_integration import (
        get_auth_url, get_token_from_code, get_user_info, 
        send_email, get_rfi_email_template, get_followup_email_template, 
        get_review_request_template, read_emails, get_email_details,
        refresh_access_token as outlook_refresh_token
    )
    from reference_api import reference_bp, init_reference_data
    from quickbooks_routes import quickbooks_bp
    from commission_routes import commission_bp
    from zoom_routes import zoom_bp
    from lead_events_routes import lead_events_bp
    from self_check_routes import self_check_bp
    from qc_routes import qc_bp
    from zoho_oauth_routes import zoho_oauth_bp
    from workdrive_routes import workdrive_bp
    from email_routes import email_bp
    from email_sending_routes import email_sending_bp
    from zoho_routes import zoho_bp
    from team_tasks_routes import team_tasks_bp
    from user_signatures_routes import user_signatures_bp
    from pipeline_routes import pipeline_bp, get_or_create_inquiry
    import zoho_integration as zoho
    import lead_events
    import event_detection
    from quickbooks_service import get_qb_service
    import email_scanner
    import token_storage
    import user_roles
    import self_check_storage
    import official_qc_storage
    import google_sheets_service
    import streaming_analyzer
    import gemini_estimate

app = Flask(__name__)
app.secret_key = Config.SECRET_KEY
CORS(app, supports_credentials=True, origins=Config.CORS_ORIGINS)

# Session cookie settings for Cloud Run
app.config['SESSION_COOKIE_SECURE'] = Config.IS_CLOUD_RUN  # HTTPS only in production
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['PERMANENT_SESSION_LIFETIME'] = 86400 * 7  # 7 days

# WorkDrive clients folder ID (needed early for cache function)
CLIENTS_FOLDER_ID = 'ar3d0456a1e2e8a4a4126894544a8e50e0735'

# WorkDrive clients cache - in-memory for same-instance reuse (short TTL)
# Primary cache is Firestore (shared across Cloud Run instances)
_workdrive_clients_cache = {
    'data': None,
    'timestamp': None,
    'ttl': 60  # 1 minute in-memory (Firestore has 5 min TTL)
}

def get_cached_clients(token):
    """Get all client folders from cache or fetch if stale.

    Cache hierarchy:
    1. In-memory cache (60s TTL) - fastest, same Cloud Run instance
    2. Firestore cache (300s TTL) - shared across all instances
    3. Zoho API fetch - slowest, updates both caches
    """
    import time
    cache = _workdrive_clients_cache

    # 1. Check in-memory cache first (fastest)
    if cache['data'] and cache['timestamp']:
        age = time.time() - cache['timestamp']
        if age < cache['ttl']:
            print(f"[WorkDrive] Using in-memory cache ({len(cache['data'])} folders, {int(age)}s old)")
            return cache['data']

    # 2. Check Firestore cache (shared across instances)
    firestore_clients = token_storage.get_workdrive_clients_cache(ttl_seconds=300)
    if firestore_clients:
        # Update in-memory cache from Firestore
        cache['data'] = firestore_clients
        cache['timestamp'] = time.time()
        print(f"[WorkDrive] Loaded {len(firestore_clients)} clients from Firestore cache")
        return firestore_clients

    # 3. Fetch from Zoho API
    print("[WorkDrive] Fetching all client folders from Zoho...")
    all_clients = []
    offset = 0
    while True:
        folders = zoho.get_folder_contents(token, CLIENTS_FOLDER_ID, limit=200, offset=offset)
        batch = [f for f in folders.get('data', []) if f.get('attributes', {}).get('is_folder')]
        if not batch:
            break
        all_clients.extend(batch)
        offset += 200
        if len(batch) < 200:
            break

    # Update both caches
    cache['data'] = all_clients
    cache['timestamp'] = time.time()
    token_storage.save_workdrive_clients_cache(all_clients)
    print(f"[WorkDrive] Cached {len(all_clients)} client folders (in-memory + Firestore)")

    return all_clients

# Print configuration status on startup
if not Config.IS_CLOUD_RUN:
    Config.print_status()

# Register blueprints
app.register_blueprint(reference_bp)
app.register_blueprint(quickbooks_bp)
app.register_blueprint(commission_bp)
app.register_blueprint(zoom_bp)
app.register_blueprint(lead_events_bp)
app.register_blueprint(self_check_bp)
app.register_blueprint(qc_bp)
app.register_blueprint(zoho_oauth_bp)
app.register_blueprint(workdrive_bp)
app.register_blueprint(email_bp)
app.register_blueprint(email_sending_bp)
app.register_blueprint(zoho_bp)
app.register_blueprint(team_tasks_bp)
app.register_blueprint(user_signatures_bp)
app.register_blueprint(pipeline_bp)
init_reference_data()  # Load reference data on startup

# Email scanner will be started on first request (see below)
# This avoids issues with gunicorn worker processes and background threads
_scanner_started = False

@app.before_request
def ensure_scanner_running():
    global _scanner_started
    if not _scanner_started:
        print("[STARTUP] Auto-starting email scanner on first request...")
        email_scanner.start_scanner()
        _scanner_started = True

# Static files directory (for Cloud Run deployment)
STATIC_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'Designer_Dashboard')
FRONTOFFICE_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'FrontOffice_Dashboard')

# Serve frontend
@app.route('/')
def serve_index():
    return send_from_directory(STATIC_DIR, 'index_modular.html')

@app.route('/frontoffice')
def serve_frontoffice():
    return send_from_directory(FRONTOFFICE_DIR, 'index.html')

@app.route('/<path:filename>')
def serve_static(filename):
    return send_from_directory(STATIC_DIR, filename)

# Use configuration from Config module
PROJECTS_BASE_PATH = Config.PROJECTS_BASE_PATH
BUILDING_CODES_PATH = Config.BUILDING_CODES_PATH
API_KEY = Config.ANTHROPIC_API_KEY


# ==================== STANDARD ERROR RESPONSES ====================

def error_response(message, code=400, details=None):
    """Create a standardized error response.
    
    Args:
        message: Human-readable error message
        code: HTTP status code (default 400)
        details: Optional additional details (dict or string)
    
    Returns:
        Tuple of (jsonify response, status code)
    """
    response = {
        'error': message,
        'status': code
    }
    if details:
        response['details'] = details
    return jsonify(response), code


def get_outlook_token():
    """Get Outlook token from session or Firestore"""
    if 'access_token' in session:
        return session['access_token']
    # Check user-specific token first
    user_email = session.get('user_email')
    if user_email:
        stored = token_storage.get_token('outlook', user_id=user_email)
        if stored and stored.get('access_token'):
            session['access_token'] = stored['access_token']
            return stored['access_token']
    # Fallback to shared token
    stored = token_storage.get_token('outlook')
    if stored and stored.get('access_token'):
        session['access_token'] = stored['access_token']
        return stored['access_token']
    return None


def get_zoho_token():
    """Get Zoho token from session or Firestore"""
    if 'zoho_access_token' in session:
        return session['zoho_access_token']
    # Check user-specific token first
    user_email = session.get('user_email')
    if user_email:
        stored = token_storage.get_token('zoho', user_id=user_email)
        if stored and stored.get('access_token'):
            session['zoho_access_token'] = stored['access_token']
            session['zoho_refresh_token'] = stored.get('refresh_token')
            return stored['access_token']
    # Fallback to shared token
    stored = token_storage.get_token('zoho')
    if stored and stored.get('access_token'):
        session['zoho_access_token'] = stored['access_token']
        session['zoho_refresh_token'] = stored.get('refresh_token')
        return stored['access_token']
    return None


# Zip code to state mapping (first 3 digits)
ZIP_TO_STATE = {
    '005': 'NY', '006': 'PR', '007': 'PR', '008': 'VI', '009': 'PR',
    '010': 'MA', '011': 'MA', '012': 'MA', '013': 'MA', '014': 'MA', '015': 'MA', '016': 'MA', '017': 'MA', '018': 'MA', '019': 'MA',
    '020': 'MA', '021': 'MA', '022': 'MA', '023': 'MA', '024': 'MA', '025': 'MA', '026': 'MA', '027': 'MA',
    '028': 'RI', '029': 'RI',
    '030': 'NH', '031': 'NH', '032': 'NH', '033': 'NH', '034': 'NH', '035': 'NH', '036': 'NH', '037': 'NH', '038': 'NH',
    '039': 'ME', '040': 'ME', '041': 'ME', '042': 'ME', '043': 'ME', '044': 'ME', '045': 'ME', '046': 'ME', '047': 'ME', '048': 'ME', '049': 'ME',
    '050': 'VT', '051': 'VT', '052': 'VT', '053': 'VT', '054': 'VT', '055': 'VT', '056': 'VT', '057': 'VT', '058': 'VT', '059': 'VT',
    '060': 'CT', '061': 'CT', '062': 'CT', '063': 'CT', '064': 'CT', '065': 'CT', '066': 'CT', '067': 'CT', '068': 'CT', '069': 'CT',
    '070': 'NJ', '071': 'NJ', '072': 'NJ', '073': 'NJ', '074': 'NJ', '075': 'NJ', '076': 'NJ', '077': 'NJ', '078': 'NJ', '079': 'NJ', '080': 'NJ', '081': 'NJ', '082': 'NJ', '083': 'NJ', '084': 'NJ', '085': 'NJ', '086': 'NJ', '087': 'NJ', '088': 'NJ', '089': 'NJ',
    '100': 'NY', '101': 'NY', '102': 'NY', '103': 'NY', '104': 'NY', '105': 'NY', '106': 'NY', '107': 'NY', '108': 'NY', '109': 'NY',
    '110': 'NY', '111': 'NY', '112': 'NY', '113': 'NY', '114': 'NY', '115': 'NY', '116': 'NY', '117': 'NY', '118': 'NY', '119': 'NY',
    '120': 'NY', '121': 'NY', '122': 'NY', '123': 'NY', '124': 'NY', '125': 'NY', '126': 'NY', '127': 'NY', '128': 'NY', '129': 'NY',
    '130': 'NY', '131': 'NY', '132': 'NY', '133': 'NY', '134': 'NY', '135': 'NY', '136': 'NY', '137': 'NY', '138': 'NY', '139': 'NY',
    '140': 'NY', '141': 'NY', '142': 'NY', '143': 'NY', '144': 'NY', '145': 'NY', '146': 'NY', '147': 'NY', '148': 'NY', '149': 'NY',
    '150': 'PA', '151': 'PA', '152': 'PA', '153': 'PA', '154': 'PA', '155': 'PA', '156': 'PA', '157': 'PA', '158': 'PA', '159': 'PA',
    '160': 'PA', '161': 'PA', '162': 'PA', '163': 'PA', '164': 'PA', '165': 'PA', '166': 'PA', '167': 'PA', '168': 'PA', '169': 'PA',
    '170': 'PA', '171': 'PA', '172': 'PA', '173': 'PA', '174': 'PA', '175': 'PA', '176': 'PA', '177': 'PA', '178': 'PA', '179': 'PA',
    '180': 'PA', '181': 'PA', '182': 'PA', '183': 'PA', '184': 'PA', '185': 'PA', '186': 'PA', '187': 'PA', '188': 'PA', '189': 'PA',
    '190': 'PA', '191': 'PA', '192': 'PA', '193': 'PA', '194': 'PA', '195': 'PA', '196': 'PA',
    '197': 'DE', '198': 'DE', '199': 'DE',
    '200': 'DC', '201': 'VA', '202': 'DC', '203': 'DC', '204': 'DC', '205': 'DC',
    '206': 'MD', '207': 'MD', '208': 'MD', '209': 'MD', '210': 'MD', '211': 'MD', '212': 'MD', '214': 'MD', '215': 'MD', '216': 'MD', '217': 'MD', '218': 'MD', '219': 'MD',
    '220': 'VA', '221': 'VA', '222': 'VA', '223': 'VA', '224': 'VA', '225': 'VA', '226': 'VA', '227': 'VA', '228': 'VA', '229': 'VA',
    '230': 'VA', '231': 'VA', '232': 'VA', '233': 'VA', '234': 'VA', '235': 'VA', '236': 'VA', '237': 'VA', '238': 'VA', '239': 'VA',
    '240': 'VA', '241': 'VA', '242': 'VA', '243': 'VA', '244': 'VA', '245': 'VA', '246': 'VA',
    '247': 'WV', '248': 'WV', '249': 'WV', '250': 'WV', '251': 'WV', '252': 'WV', '253': 'WV', '254': 'WV', '255': 'WV', '256': 'WV', '257': 'WV', '258': 'WV', '259': 'WV', '260': 'WV', '261': 'WV', '262': 'WV', '263': 'WV', '264': 'WV', '265': 'WV', '266': 'WV', '267': 'WV', '268': 'WV',
    '270': 'NC', '271': 'NC', '272': 'NC', '273': 'NC', '274': 'NC', '275': 'NC', '276': 'NC', '277': 'NC', '278': 'NC', '279': 'NC',
    '280': 'NC', '281': 'NC', '282': 'NC', '283': 'NC', '284': 'NC', '285': 'NC', '286': 'NC', '287': 'NC', '288': 'NC', '289': 'NC',
    '290': 'SC', '291': 'SC', '292': 'SC', '293': 'SC', '294': 'SC', '295': 'SC', '296': 'SC', '297': 'SC', '298': 'SC', '299': 'SC',
    '300': 'GA', '301': 'GA', '302': 'GA', '303': 'GA', '304': 'GA', '305': 'GA', '306': 'GA', '307': 'GA', '308': 'GA', '309': 'GA',
    '310': 'GA', '311': 'GA', '312': 'GA', '313': 'GA', '314': 'GA', '315': 'GA', '316': 'GA', '317': 'GA', '318': 'GA', '319': 'GA',
    '320': 'FL', '321': 'FL', '322': 'FL', '323': 'FL', '324': 'FL', '325': 'FL', '326': 'FL', '327': 'FL', '328': 'FL', '329': 'FL',
    '330': 'FL', '331': 'FL', '332': 'FL', '333': 'FL', '334': 'FL', '335': 'FL', '336': 'FL', '337': 'FL', '338': 'FL', '339': 'FL',
    '340': 'FL', '341': 'FL', '342': 'FL', '344': 'FL', '346': 'FL', '347': 'FL', '349': 'FL',
    '350': 'AL', '351': 'AL', '352': 'AL', '354': 'AL', '355': 'AL', '356': 'AL', '357': 'AL', '358': 'AL', '359': 'AL',
    '360': 'AL', '361': 'AL', '362': 'AL', '363': 'AL', '364': 'AL', '365': 'AL', '366': 'AL', '367': 'AL', '368': 'AL', '369': 'AL',
    '370': 'TN', '371': 'TN', '372': 'TN', '373': 'TN', '374': 'TN', '375': 'TN', '376': 'TN', '377': 'TN', '378': 'TN', '379': 'TN',
    '380': 'TN', '381': 'TN', '382': 'TN', '383': 'TN', '384': 'TN', '385': 'TN',
    '386': 'MS', '387': 'MS', '388': 'MS', '389': 'MS', '390': 'MS', '391': 'MS', '392': 'MS', '393': 'MS', '394': 'MS', '395': 'MS', '396': 'MS', '397': 'MS',
    '400': 'KY', '401': 'KY', '402': 'KY', '403': 'KY', '404': 'KY', '405': 'KY', '406': 'KY', '407': 'KY', '408': 'KY', '409': 'KY',
    '410': 'KY', '411': 'KY', '412': 'KY', '413': 'KY', '414': 'KY', '415': 'KY', '416': 'KY', '417': 'KY', '418': 'KY',
    '420': 'KY', '421': 'KY', '422': 'KY', '423': 'KY', '424': 'KY', '425': 'KY', '426': 'KY', '427': 'KY',
    '430': 'OH', '431': 'OH', '432': 'OH', '433': 'OH', '434': 'OH', '435': 'OH', '436': 'OH', '437': 'OH', '438': 'OH', '439': 'OH',
    '440': 'OH', '441': 'OH', '442': 'OH', '443': 'OH', '444': 'OH', '445': 'OH', '446': 'OH', '447': 'OH', '448': 'OH', '449': 'OH',
    '450': 'OH', '451': 'OH', '452': 'OH', '453': 'OH', '454': 'OH', '455': 'OH', '456': 'OH', '457': 'OH', '458': 'OH', '459': 'OH',
    '460': 'IN', '461': 'IN', '462': 'IN', '463': 'IN', '464': 'IN', '465': 'IN', '466': 'IN', '467': 'IN', '468': 'IN', '469': 'IN',
    '470': 'IN', '471': 'IN', '472': 'IN', '473': 'IN', '474': 'IN', '475': 'IN', '476': 'IN', '477': 'IN', '478': 'IN', '479': 'IN',
    '480': 'MI', '481': 'MI', '482': 'MI', '483': 'MI', '484': 'MI', '485': 'MI', '486': 'MI', '487': 'MI', '488': 'MI', '489': 'MI',
    '490': 'MI', '491': 'MI', '492': 'MI', '493': 'MI', '494': 'MI', '495': 'MI', '496': 'MI', '497': 'MI', '498': 'MI', '499': 'MI',
    '500': 'IA', '501': 'IA', '502': 'IA', '503': 'IA', '504': 'IA', '505': 'IA', '506': 'IA', '507': 'IA', '508': 'IA', '509': 'IA',
    '510': 'IA', '511': 'IA', '512': 'IA', '513': 'IA', '514': 'IA', '515': 'IA', '516': 'IA', '520': 'IA', '521': 'IA', '522': 'IA', '523': 'IA', '524': 'IA', '525': 'IA', '526': 'IA', '527': 'IA', '528': 'IA',
    '530': 'WI', '531': 'WI', '532': 'WI', '534': 'WI', '535': 'WI', '537': 'WI', '538': 'WI', '539': 'WI',
    '540': 'WI', '541': 'WI', '542': 'WI', '543': 'WI', '544': 'WI', '545': 'WI', '546': 'WI', '547': 'WI', '548': 'WI', '549': 'WI',
    '550': 'MN', '551': 'MN', '553': 'MN', '554': 'MN', '555': 'MN', '556': 'MN', '557': 'MN', '558': 'MN', '559': 'MN',
    '560': 'MN', '561': 'MN', '562': 'MN', '563': 'MN', '564': 'MN', '565': 'MN', '566': 'MN', '567': 'MN',
    '570': 'SD', '571': 'SD', '572': 'SD', '573': 'SD', '574': 'SD', '575': 'SD', '576': 'SD', '577': 'SD',
    '580': 'ND', '581': 'ND', '582': 'ND', '583': 'ND', '584': 'ND', '585': 'ND', '586': 'ND', '587': 'ND', '588': 'ND',
    '590': 'MT', '591': 'MT', '592': 'MT', '593': 'MT', '594': 'MT', '595': 'MT', '596': 'MT', '597': 'MT', '598': 'MT', '599': 'MT',
    '600': 'IL', '601': 'IL', '602': 'IL', '603': 'IL', '604': 'IL', '605': 'IL', '606': 'IL', '607': 'IL', '608': 'IL', '609': 'IL',
    '610': 'IL', '611': 'IL', '612': 'IL', '613': 'IL', '614': 'IL', '615': 'IL', '616': 'IL', '617': 'IL', '618': 'IL', '619': 'IL',
    '620': 'IL', '622': 'IL', '623': 'IL', '624': 'IL', '625': 'IL', '626': 'IL', '627': 'IL', '628': 'IL', '629': 'IL',
    '630': 'MO', '631': 'MO', '633': 'MO', '634': 'MO', '635': 'MO', '636': 'MO', '637': 'MO', '638': 'MO', '639': 'MO',
    '640': 'MO', '641': 'MO', '644': 'MO', '645': 'MO', '646': 'MO', '647': 'MO', '648': 'MO', '649': 'MO',
    '650': 'MO', '651': 'MO', '652': 'MO', '653': 'MO', '654': 'MO', '655': 'MO', '656': 'MO', '657': 'MO', '658': 'MO',
    '660': 'KS', '661': 'KS', '662': 'KS', '664': 'KS', '665': 'KS', '666': 'KS', '667': 'KS', '668': 'KS', '669': 'KS',
    '670': 'KS', '671': 'KS', '672': 'KS', '673': 'KS', '674': 'KS', '675': 'KS', '676': 'KS', '677': 'KS', '678': 'KS', '679': 'KS',
    '680': 'NE', '681': 'NE', '683': 'NE', '684': 'NE', '685': 'NE', '686': 'NE', '687': 'NE', '688': 'NE', '689': 'NE',
    '690': 'NE', '691': 'NE', '692': 'NE', '693': 'NE',
    '700': 'LA', '701': 'LA', '703': 'LA', '704': 'LA', '705': 'LA', '706': 'LA', '707': 'LA', '708': 'LA',
    '710': 'LA', '711': 'LA', '712': 'LA', '713': 'LA', '714': 'LA',
    '716': 'AR', '717': 'AR', '718': 'AR', '719': 'AR', '720': 'AR', '721': 'AR', '722': 'AR', '723': 'AR', '724': 'AR', '725': 'AR', '726': 'AR', '727': 'AR', '728': 'AR', '729': 'AR',
    '730': 'OK', '731': 'OK', '734': 'OK', '735': 'OK', '736': 'OK', '737': 'OK', '738': 'OK', '739': 'OK',
    '740': 'OK', '741': 'OK', '743': 'OK', '744': 'OK', '745': 'OK', '746': 'OK', '747': 'OK', '748': 'OK', '749': 'OK',
    '750': 'TX', '751': 'TX', '752': 'TX', '753': 'TX', '754': 'TX', '755': 'TX', '756': 'TX', '757': 'TX', '758': 'TX', '759': 'TX',
    '760': 'TX', '761': 'TX', '762': 'TX', '763': 'TX', '764': 'TX', '765': 'TX', '766': 'TX', '767': 'TX', '768': 'TX', '769': 'TX',
    '770': 'TX', '771': 'TX', '772': 'TX', '773': 'TX', '774': 'TX', '775': 'TX', '776': 'TX', '777': 'TX', '778': 'TX', '779': 'TX',
    '780': 'TX', '781': 'TX', '782': 'TX', '783': 'TX', '784': 'TX', '785': 'TX', '786': 'TX', '787': 'TX', '788': 'TX', '789': 'TX',
    '790': 'TX', '791': 'TX', '792': 'TX', '793': 'TX', '794': 'TX', '795': 'TX', '796': 'TX', '797': 'TX', '798': 'TX', '799': 'TX',
    '800': 'CO', '801': 'CO', '802': 'CO', '803': 'CO', '804': 'CO', '805': 'CO', '806': 'CO', '807': 'CO', '808': 'CO', '809': 'CO',
    '810': 'CO', '811': 'CO', '812': 'CO', '813': 'CO', '814': 'CO', '815': 'CO', '816': 'CO',
    '820': 'WY', '821': 'WY', '822': 'WY', '823': 'WY', '824': 'WY', '825': 'WY', '826': 'WY', '827': 'WY', '828': 'WY', '829': 'WY', '830': 'WY', '831': 'WY',
    '832': 'ID', '833': 'ID', '834': 'ID', '835': 'ID', '836': 'ID', '837': 'ID', '838': 'ID',
    '840': 'UT', '841': 'UT', '842': 'UT', '843': 'UT', '844': 'UT', '845': 'UT', '846': 'UT', '847': 'UT',
    '850': 'AZ', '851': 'AZ', '852': 'AZ', '853': 'AZ', '855': 'AZ', '856': 'AZ', '857': 'AZ', '859': 'AZ', '860': 'AZ', '863': 'AZ', '864': 'AZ', '865': 'AZ',
    '870': 'NM', '871': 'NM', '872': 'NM', '873': 'NM', '874': 'NM', '875': 'NM', '877': 'NM', '878': 'NM', '879': 'NM',
    '880': 'NM', '881': 'NM', '882': 'NM', '883': 'NM', '884': 'NM',
    '889': 'NV', '890': 'NV', '891': 'NV', '893': 'NV', '894': 'NV', '895': 'NV', '897': 'NV', '898': 'NV',
    '900': 'CA', '901': 'CA', '902': 'CA', '903': 'CA', '904': 'CA', '905': 'CA', '906': 'CA', '907': 'CA', '908': 'CA', '909': 'CA',
    '910': 'CA', '911': 'CA', '912': 'CA', '913': 'CA', '914': 'CA', '915': 'CA', '916': 'CA', '917': 'CA', '918': 'CA',
    '920': 'CA', '921': 'CA', '922': 'CA', '923': 'CA', '924': 'CA', '925': 'CA', '926': 'CA', '927': 'CA', '928': 'CA',
    '930': 'CA', '931': 'CA', '932': 'CA', '933': 'CA', '934': 'CA', '935': 'CA', '936': 'CA', '937': 'CA', '938': 'CA', '939': 'CA',
    '940': 'CA', '941': 'CA', '942': 'CA', '943': 'CA', '944': 'CA', '945': 'CA', '946': 'CA', '947': 'CA', '948': 'CA', '949': 'CA',
    '950': 'CA', '951': 'CA', '952': 'CA', '953': 'CA', '954': 'CA', '955': 'CA', '956': 'CA', '957': 'CA', '958': 'CA', '959': 'CA',
    '960': 'CA', '961': 'CA',
    '967': 'HI', '968': 'HI',
    '970': 'OR', '971': 'OR', '972': 'OR', '973': 'OR', '974': 'OR', '975': 'OR', '976': 'OR', '977': 'OR', '978': 'OR', '979': 'OR',
    '980': 'WA', '981': 'WA', '982': 'WA', '983': 'WA', '984': 'WA', '985': 'WA', '986': 'WA', '988': 'WA', '989': 'WA',
    '990': 'WA', '991': 'WA', '992': 'WA', '993': 'WA', '994': 'WA',
    '995': 'AK', '996': 'AK', '997': 'AK', '998': 'AK', '999': 'AK'
}

@app.route('/api/projects', methods=['GET'])
def get_projects():
    """Get list of all projects"""
    try:
        projects = []
        base_path = Path(PROJECTS_BASE_PATH)
        
        if not base_path.exists():
            return jsonify({'projects': [], 'error': 'Projects folder not found'})
        
        for item in base_path.iterdir():
            if item.is_dir() and not item.name.startswith('.'):
                projects.append({
                    'id': item.name,
                    'name': item.name,
                    'path': str(item),
                    'status': 'New Assignment'
                })
        
        return jsonify({'projects': projects})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/projects/<project_id>/analyze', methods=['POST'])
def analyze_project(project_id):
    """Analyze a project and return summary + RFI"""
    try:
        project_path = Path(PROJECTS_BASE_PATH) / project_id
        
        if not project_path.exists():
            return error_response('Project not found', 404)
        
        # Run analyzer for all projects
        analyzer = ProjectAnalyzer(str(project_path))
        analyzer.scan_project_files()
        analyzer.extract_all_content()
        results = analyzer.generate_ai_summary(API_KEY)
        
        # Parse the AI response into structured format
        parsed = parse_ai_analysis(results['ai_analysis'])
        
        return jsonify({
            'project_name': project_id,
            'summary': parsed['project_summary'],
            'missing_info': parsed['missing_info'],
            'files_provided': parsed['files_provided'],
            'file_count': {
                'rup': len(analyzer.project_data['rup_files']),
                'pdf': len(analyzer.project_data['pdf_files']),
                'emails': len(analyzer.project_data['emails'])
            }
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/projects/<project_id>/generate-rfi', methods=['POST'])
def generate_rfi(project_id):
    """Generate RFI email from selected missing items"""
    try:
        data = request.json
        project_summary = data.get('summary', {})
        selected_items = data.get('selected_items', [])
        
        rfi_email = generate_rfi_email(project_summary, selected_items)
        
        return jsonify(rfi_email)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/projects/<project_id>/analysis', methods=['GET'])
def get_existing_analysis(project_id):
    """Get existing analysis if available"""
    try:
        project_path = Path(PROJECTS_BASE_PATH) / project_id
        analysis_file = project_path / 'AI_Analysis_Results.json'
        
        if analysis_file.exists():
            import json
            with open(analysis_file, 'r') as f:
                data = json.load(f)
            return jsonify(data)
        else:
            return jsonify({'error': 'No analysis found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/workdrive/analyze-project', methods=['POST'])
def analyze_workdrive_project():
    """Analyze a project from WorkDrive files - uses standard folder structure"""
    import anthropic
    import PyPDF2
    import io
    
    token = get_zoho_token()
    if not token:
        return jsonify({'error': 'Not connected to Zoho'}), 401
    
    try:
        data = request.json
        project_name = data.get('project_name')
        client_name = data.get('client_name')
        
        if not project_name or not client_name:
            return jsonify({'error': 'project_name and client_name required'}), 400
        
        print(f"[ANALYZE] Looking for client folder: '{client_name}'")
        
        # Find client folder using cache
        client_folder = None
        client_name_lower = client_name.lower()
        cached_clients = get_cached_clients(token)
        
        for folder in cached_clients:
            name = folder.get('attributes', {}).get('name', '')
            if name.lower() == client_name_lower:
                client_folder = folder
                print(f"[ANALYZE] Found client folder in cache: {name}")
                break
        
        if not client_folder:
            print(f"[ANALYZE] Exact match not found, trying fuzzy search in cache...")
            # Fuzzy search using cache
            similar = []
            search_lower = client_name.lower()
            for folder in cached_clients:
                name = folder.get('attributes', {}).get('name', '')
                name_lower = name.lower()
                if search_lower in name_lower or name_lower in search_lower:
                    similar.append(name)
            
            if similar:
                return jsonify({
                    'error': f'Client folder not found: {client_name}',
                    'similar_folders': similar[:10],
                    'suggestion': f'Did you mean one of these? {", ".join(similar[:5])}'
                }), 404
            else:
                return jsonify({'error': f'Client folder not found: {client_name}'}), 404
        
        client_folder_id = client_folder.get('id')
        print(f"[ANALYZE] Found client folder: {client_folder_id}")
        
        # Get project files using the standard folder structure
        project_files = zoho.get_project_files_from_workdrive(token, client_folder_id, project_name)
        
        # Count what we found
        total_files = sum([
            len(project_files.get('forms', [])),
            len(project_files.get('files_from_client', [])),
            len(project_files.get('working_drawings', [])),
            len(project_files.get('emails', []))
        ])
        
        print(f"[ANALYZE] Total files found: {total_files}")
        
        if total_files == 0:
            return jsonify({
                'project_name': project_name,
                'summary': {
                    'project_name': project_name,
                    'client_name': client_name
                },
                'missing_info': [{'item': 'No files found', 'reason': 'Project folders are empty'}],
                'files_provided': [],
                'file_count': {
                    'forms': 0, 'files_from_client': 0, 'working_drawings': 0, 'emails': 0
                }
            })
        
        # Download files from forms, files_from_client, emails (not working_drawings - too large)
        files_content = []
        files_downloaded = []
        downloadable_extensions = ['.pdf', '.txt', '.msg', '.eml']
        
        for folder_key in ['forms', 'files_from_client', 'emails']:
            for f in project_files.get(folder_key, []):
                fname = f.get('attributes', {}).get('name', 'Unknown')
                file_id = f.get('id')
                
                if any(fname.lower().endswith(e) for e in downloadable_extensions):
                    print(f"[ANALYZE] Downloading {folder_key}/{fname}...")
                    try:
                        content = zoho.download_file_content(token, file_id)
                        if content:
                            files_content.append({
                                'name': fname,
                                'folder': folder_key,
                                'content': content
                            })
                            files_downloaded.append(fname)
                            print(f"[ANALYZE] Downloaded {len(content)} bytes")
                    except Exception as e:
                        print(f"[ANALYZE] ERROR downloading {fname}: {e}")
        
        print(f"[ANALYZE] Downloaded {len(files_downloaded)} files")
        
        if not files_content:
            return jsonify({
                'project_name': project_name,
                'summary': {'project_name': project_name, 'client_name': client_name},
                'missing_info': [{'item': 'No downloadable files', 'reason': 'Could not download files'}],
                'files_provided': [],
                'file_count': {
                    'forms': len(project_files.get('forms', [])),
                    'files_from_client': len(project_files.get('files_from_client', [])),
                    'working_drawings': len(project_files.get('working_drawings', [])),
                    'emails': len(project_files.get('emails', []))
                }
            })
        
        # Extract text from downloaded files
        extracted_texts = []
        for file_info in files_content:
            fname = file_info['name']
            content = file_info['content']
            folder = file_info['folder']
            text = ""
            
            if fname.lower().endswith('.pdf'):
                try:
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                    
                    # First, extract fillable form field values
                    fields = pdf_reader.get_fields()
                    if fields:
                        text += "=== FORM FIELD VALUES ===\n"
                        for field_name, field_data in fields.items():
                            value = field_data.get('/V', '')
                            # Handle different value types
                            if isinstance(value, str) and value and value not in ['', '/Off']:
                                text += f"{field_name}: {value}\n"
                            elif hasattr(value, 'get_object'):
                                try:
                                    resolved = str(value.get_object())
                                    if resolved and resolved not in ['', '/Off']:
                                        text += f"{field_name}: {resolved}\n"
                                except:
                                    pass
                        text += "=== END FORM FIELDS ===\n\n"
                    
                    # Then extract page text
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    print(f"[ANALYZE] Extracted {len(text)} chars from {fname} (fields: {len(fields) if fields else 0})")
                except Exception as e:
                    print(f"[ANALYZE] PDF extraction failed for {fname}: {e}")
            elif fname.lower().endswith(('.txt', '.msg', '.eml')):
                try:
                    text = content.decode('utf-8', errors='ignore')
                except:
                    text = str(content)
            
            if text:
                extracted_texts.append({
                    'filename': fname,
                    'folder': folder,
                    'text': text[:15000]
                })
        
        # Build AI prompt
        all_text = ""
        for doc in extracted_texts:
            all_text += f"\n\n=== FILE: {doc['filename']} (from {doc['folder']}) ===\n{doc['text']}"
        
        ai_prompt = f"""Analyze these HVAC project documents and extract ALL information from Summary Forms and Estimates.

PROJECT: {project_name}
CLIENT: {client_name}

DOCUMENTS:
{all_text[:30000]}

IMPORTANT: Look carefully at Summary Form PDFs for equipment details like:
- Number of systems (# Systems field)
- Cooling type (Heat Pump Split System, Mini-Split, etc.)
- Heating type (Heat Pump, Furnace, etc.)
- Brand (Carrier, Trane, etc.)
- SEER rating
- Phase/Volts
- Wall types, R-values, insulation details

Provide a JSON response with these exact field names:
{{
    "project_summary": {{
        "project_name": "{project_name}",
        "client_name": "{client_name}",
        "location": "full address if found",
        "project_type": "RNC/RREN/CNC/CREN/Bare RNC or N/A",
        "design_service": "Full Design/Basic/Bare/Manual J Only or N/A",
        "project_cost": "dollar amount from estimate or N/A",
        "total_area_sqft": "square footage or N/A",
        "system_type": "e.g. Heat Pump Split System, Mini-Split, Furnace+AC or N/A",
        "num_systems": "number from # Systems field or N/A",
        "fuel_type": "gas/electric/propane or N/A",
        "construction_type": "new/existing/renovation or N/A",
        "brand": "equipment brand if specified or N/A",
        "seer": "SEER rating if specified or N/A"
    }},
    "missing_info": [
        {{"item": "field name", "reason": "why it's missing or needed for design"}}
    ],
    "files_provided": ["list of analyzed files"],
    "notes": "important observations for the designer"
}}

Return ONLY valid JSON, no markdown or extra text."""

        print(f"[ANALYZE] Sending to AI ({len(ai_prompt)} chars)...")
        
        client = anthropic.Anthropic(api_key=Config.ANTHROPIC_API_KEY)
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2000,
            messages=[{"role": "user", "content": ai_prompt}]
        )
        
        ai_response = response.content[0].text.strip()
        print(f"[ANALYZE] AI response received")
        
        # Parse JSON
        import re
        if ai_response.startswith('```'):
            ai_response = ai_response.split('```')[1]
            if ai_response.startswith('json'):
                ai_response = ai_response[4:]
        
        try:
            json_match = re.search(r'\{[\s\S]*\}', ai_response)
            if json_match:
                result = json.loads(json_match.group())
                # Rename project_summary to summary for frontend compatibility
                if 'project_summary' in result:
                    result['summary'] = result.pop('project_summary')
            else:
                result = {'error': 'Could not parse AI response'}
        except json.JSONDecodeError as e:
            print(f"[ANALYZE] JSON parse error: {e}")
            result = {'summary': {'project_name': project_name, 'client_name': client_name}}
        
        result['files_analyzed'] = files_downloaded
        result['file_count'] = {
            'forms': len(project_files.get('forms', [])),
            'files_from_client': len(project_files.get('files_from_client', [])),
            'working_drawings': len(project_files.get('working_drawings', [])),
            'emails': len(project_files.get('emails', []))
        }
        
        return jsonify(result)
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/workdrive/analyze-project-stream/init', methods=['POST'])
def init_streaming_analysis():
    """Initialize streaming analysis - returns file count and session ID"""
    token = get_zoho_token()
    if not token:
        return jsonify({'error': 'Not connected to Zoho'}), 401
    
    try:
        data = request.json
        project_name = data.get('project_name')
        client_name = data.get('client_name')
        
        if not project_name or not client_name:
            return jsonify({'error': 'project_name and client_name required'}), 400
        
        cached_clients = get_cached_clients(token)
        result, error = streaming_analyzer.get_project_file_count(
            token, client_name, project_name, cached_clients
        )
        
        if error:
            return jsonify({'error': error}), 404
        
        # Generate unique session ID and store in Firestore
        import uuid
        stream_session_id = str(uuid.uuid4())
        user_email = session.get('user_email', 'anonymous')
        
        token_storage.save_session_data(stream_session_id, {
            'files': result['files'],
            'project_name': project_name,
            'client_name': client_name,
            'user_email': user_email
        })
        
        return jsonify({
            'session_id': stream_session_id,
            'total_files': result['total_count'],
            'file_count': result['file_count'],
            'files': [f['name'] for f in result['files']]
        })
        
    except Exception as e:
        import traceback
        import uuid
        error_id = str(uuid.uuid4())[:8].upper()
        traceback.print_exc()
        print(f"[ERROR {error_id}] init_streaming_analysis: {e}")
        return jsonify({
            'error': f"Something went wrong on our end. Error ID: {error_id}. Let Tom know if this keeps happening.",
            'error_id': error_id
        }), 500


@app.route('/api/workdrive/analyze-project-stream/file/<int:index>', methods=['POST'])
def analyze_single_file_endpoint(index):
    """Analyze a single file by index"""
    token = get_zoho_token()
    if not token:
        return jsonify({'error': 'Not connected to Zoho'}), 401
    
    try:
        data = request.json or {}
        stream_session_id = data.get('session_id')
        
        if not stream_session_id:
            return jsonify({'error': 'session_id required'}), 400
        
        # Get session data from Firestore
        stream_data = token_storage.get_session_data(stream_session_id)
        if not stream_data:
            return jsonify({'error': 'Invalid or expired session'}), 400
        
        files = stream_data.get('files', [])
        project_name = stream_data.get('project_name')
        client_name = stream_data.get('client_name')
        
        if not files or index >= len(files):
            return jsonify({'error': 'Invalid file index'}), 400
        
        file_item = files[index]
        file_info = file_item['file_info']
        folder_key = file_item['folder']
        
        # Download and extract text
        file_data = streaming_analyzer.analyze_single_file(token, file_info, folder_key)
        
        if not file_data:
            return jsonify({
                'index': index,
                'filename': file_item['name'],
                'status': 'skipped',
                'analysis': {}
            })
        
        # Analyze with AI
        client = anthropic.Anthropic(api_key=Config.ANTHROPIC_API_KEY)
        analysis = streaming_analyzer.analyze_file_with_ai(
            client, file_data, project_name, client_name
        )
        
        return jsonify({
            'index': index,
            'filename': file_item['name'],
            'status': 'complete',
            'analysis': analysis
        })
        
    except Exception as e:
        import traceback
        import uuid
        error_id = str(uuid.uuid4())[:8].upper()
        traceback.print_exc()
        print(f"[ERROR {error_id}] analyze_single_file: {e}")
        # Return skipped status instead of failing - let analysis continue
        return jsonify({
            'index': index,
            'filename': file_item.get('name', 'Unknown') if 'file_item' in dir() else 'Unknown',
            'status': 'skipped',
            'error': f"Couldn't read this file. Error ID: {error_id}",
            'analysis': {}
        })


@app.route('/api/workdrive/analyze-project-stream/finalize', methods=['POST'])
def finalize_streaming_analysis():
    """Combine all file analyses into final summary"""
    try:
        data = request.json
        file_analyses = data.get('file_analyses', [])
        stream_session_id = data.get('session_id')
        
        if not stream_session_id:
            return jsonify({'error': 'session_id required'}), 400
        
        # Get session data from Firestore
        stream_data = token_storage.get_session_data(stream_session_id)
        if not stream_data:
            return jsonify({'error': 'Invalid or expired session'}), 400
        
        project_name = stream_data.get('project_name')
        client_name = stream_data.get('client_name')
        
        client = anthropic.Anthropic(api_key=Config.ANTHROPIC_API_KEY)
        result = streaming_analyzer.combine_file_analyses(
            client, file_analyses, project_name, client_name
        )
        
        result['files_analyzed'] = [a.get('filename') for a in file_analyses if a.get('status') == 'complete']
        
        # Clean up Firestore session
        token_storage.delete_session_data(stream_session_id)
        
        return jsonify(result)
        
    except Exception as e:
        import traceback
        import uuid
        error_id = str(uuid.uuid4())[:8].upper()
        traceback.print_exc()
        print(f"[ERROR {error_id}] finalize_streaming_analysis: {e}")
        return jsonify({
            'error': f"Something went wrong creating the summary. Error ID: {error_id}. Let Tom know if this keeps happening.",
            'error_id': error_id
        }), 500


@app.route('/api/building-codes/<zip_code>', methods=['GET'])
def get_building_codes(zip_code):
    """Get building codes for a zip code. Pass ?type=commercial for commercial codes."""
    try:
        import json
        
        # Check if commercial type requested
        building_type = request.args.get('type', 'residential').lower()
        is_commercial = building_type == 'commercial'
        
        # Get state from zip code prefix
        zip_prefix = zip_code[:3]
        state_abbr = ZIP_TO_STATE.get(zip_prefix)
        
        if not state_abbr:
            return jsonify({'error': f'Unknown zip code prefix: {zip_prefix}'}), 404
        
        # Load state building codes - commercial or residential
        if is_commercial:
            state_file = Path(BUILDING_CODES_PATH) / f'{state_abbr}_commercial.json'
            # Fall back to residential if commercial doesn't exist
            if not state_file.exists():
                state_file = Path(BUILDING_CODES_PATH) / f'{state_abbr}.json'
                is_commercial = False
        else:
            state_file = Path(BUILDING_CODES_PATH) / f'{state_abbr}.json'
        
        if not state_file.exists():
            # Return demo data if file doesn't exist
            return jsonify({
                'state': state_abbr,
                'state_abbr': state_abbr,
                'zip_code': zip_code,
                'schema': 'commercial' if is_commercial else 'residential',
                'message': f'Building codes for {state_abbr} not yet loaded. Demo mode.',
                'energy_code': 'Check local jurisdiction',
                'hvac': {
                    'manual_j_required': True,
                    'manual_s_required': True
                },
                'is_demo': True
            })
        
        with open(state_file, 'r') as f:
            codes = json.load(f)
        
        codes['state_abbr'] = state_abbr
        codes['zip_code'] = zip_code
        
        return jsonify(codes)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ==================== OUTLOOK AUTHENTICATION ====================

@app.route('/auth/login')
def auth_login():
    """Redirect to Microsoft login - fresh OAuth only"""
    # Clear any existing session to force fresh login
    user_email = session.get('user_email')
    if user_email:
        print(f"[AUTH] Clearing session for {user_email} before new login")
    session.clear()
    
    # Store return URL for after callback
    return_url = request.args.get('return_url', '/')
    session['outlook_return_url'] = return_url
    auth_url = get_auth_url()
    return redirect(auth_url)

@app.route('/auth/callback')
def auth_callback():
    """Handle OAuth callback from Microsoft - Outlook only, no Zoho chain"""
    auth_code = request.args.get('code')
    print(f"[AUTH] Callback received, code starts with: {auth_code[:20] if auth_code else 'None'}...")
    
    # Get return URL from session (default to /)
    return_url = session.pop('outlook_return_url', '/')
    
    if not auth_code:
        return jsonify({'error': 'No auth code received'}), 400
    
    # Check if we already processed this code (prevent double-submission)
    if session.get('last_auth_code') == auth_code[:50]:
        print("[AUTH] Code already processed, redirecting to dashboard")
        return redirect(return_url)
    
    result = get_token_from_code(auth_code)
    print(f"[AUTH] Token result: {'success' if 'access_token' in result else result.get('error', 'unknown error')}")
    
    if 'access_token' in result:
        session.permanent = True
        session['access_token'] = result['access_token']
        session['last_auth_code'] = auth_code[:50]
        
        user = get_user_info(result['access_token'])
        user_email = user.get('mail') or user.get('userPrincipalName')
        session['user_email'] = user_email
        session['user_name'] = user.get('displayName')
        session.modified = True
        
        # Save Outlook token per-user
        token_storage.save_token(
            service='outlook',
            access_token=result['access_token'],
            refresh_token=result.get('refresh_token'),
            updated_by=user_email,
            user_id=user_email
        )
        print(f"[AUTH] Saved Outlook token for {user_email}")
        
        # Also save as shared mailbox token if this is tom@procalcs.net
        if user_email and user_email.lower() == 'tom@procalcs.net':
            token_storage.save_shared_mailbox_token(
                mailbox_email='tom@procalcs.net',
                service='outlook',
                access_token=result['access_token'],
                refresh_token=result.get('refresh_token'),
                updated_by=user_email
            )
            print(f"[AUTH] Saved shared mailbox token for tom@procalcs.net")
        
        # Try to restore Zoho from stored token (per-user ONLY, no shared fallback)
        zoho_connected = False
        stored_zoho = token_storage.get_token('zoho', user_id=user_email)
        if stored_zoho and stored_zoho.get('refresh_token'):
            try:
                zoho_result = zoho.refresh_access_token(stored_zoho['refresh_token'])
                if zoho_result.get('access_token'):
                    test_result = zoho.get_portals(zoho_result['access_token'])
                    if test_result.get('portals') or test_result.get('code') == 0:
                        session['zoho_access_token'] = zoho_result['access_token']
                        session['zoho_refresh_token'] = zoho_result.get('refresh_token', stored_zoho['refresh_token'])
                        token_storage.save_token(
                            service='zoho',
                            access_token=zoho_result['access_token'],
                            refresh_token=zoho_result.get('refresh_token', stored_zoho['refresh_token']),
                            updated_by=user_email,
                            user_id=user_email
                        )
                        zoho_connected = True
                        print(f"[AUTH] Auto-restored Zoho for {user_email}")
            except Exception as e:
                print(f"[AUTH] Zoho auto-restore failed (user will need to reconnect): {e}")
        
        # NO AUTO-CHAIN TO ZOHO - user connects Zoho separately if needed
        
        # Build redirect URL
        from urllib.parse import urlencode, urlparse, parse_qs, urlunparse
        parsed = urlparse(return_url)
        query_params = parse_qs(parsed.query)
        query_params['auth'] = ['success']
        query_params['email'] = [user_email]
        if zoho_connected:
            query_params['zoho'] = ['connected']
        new_query = urlencode(query_params, doseq=True)
        redirect_url = urlunparse((parsed.scheme, parsed.netloc, parsed.path, parsed.params, new_query, parsed.fragment))
        
        # Set cookie for user identification
        response = make_response(redirect(redirect_url))
        response.set_cookie(
            'procalcs_uid',
            user_email,
            httponly=True,
            secure=Config.IS_CLOUD_RUN,
            samesite='Lax',
            path='/',
            max_age=30*24*60*60  # 30 days
        )
        print(f"[AUTH] Login complete for {user_email}, zoho_connected={zoho_connected}")
        return response
    else:
        return jsonify({'error': 'Failed to get token', 'details': result}), 400

@app.route('/auth/status')
def auth_status():
    """Production-safe auth restore with validation caching.

    Performance optimization: Caches validation results in session to avoid
    repeated Graph/Zoho API calls on every page load. Cache TTL is 5 minutes.

    Flow:
    1) Check session cache - if recently validated, return cached status
    2) Validate session Outlook token with Graph
    3) If missing/invalid, restore from Firestore (per-user via ?email=, then shared)
    4) If Outlook valid, restore Zoho from Firestore (per-user then shared)
    5) Cache validation results in session
    """
    import time

    # Check if we have a recent validation cache (5 min TTL)
    VALIDATION_TTL = 300  # 5 minutes
    cache_time = session.get("_auth_validated_at")
    force_revalidate = request.args.get("force") == "true"

    if not force_revalidate and cache_time and session.get("access_token"):
        age = time.time() - cache_time
        if age < VALIDATION_TTL:
            # Return cached status without re-validating
            email = session.get("user_email")
            name = session.get("user_name")
            zoho_connected = bool(session.get("zoho_access_token"))
            shared_mailbox_connected = session.get("_shared_mailbox_valid", False)

            if email:
                permissions = user_roles.get_user_permissions(email)
                print(f"[AUTH] Using cached status for {email} ({int(age)}s old)")
                return jsonify({
                    "authenticated": True,
                    "email": email,
                    "name": name,
                    "role": permissions["role"],
                    "allowed_tabs": permissions["allowed_tabs"],
                    "zoho_connected": zoho_connected,
                    "shared_mailbox_connected": shared_mailbox_connected,
                    "cached": True
                })

    def _graph_me(access_token: str):
        """Return (ok:bool, user_json:dict|None, status:int, text:str)."""
        try:
            headers = {"Authorization": f"Bearer {access_token}"}
            resp = requests.get("https://graph.microsoft.com/v1.0/me", headers=headers, timeout=7)
            if resp.status_code == 200:
                return True, resp.json(), resp.status_code, ""
            return False, None, resp.status_code, resp.text[:300]
        except Exception as e:
            return False, None, 0, str(e)

    def _restore_outlook_from_firestore(preferred_email):
        """
        Try per-user token first (preferred_email), then shared.
        If access token is expired but refresh token exists, refresh it.
        Returns (ok, email, name)
        """
        candidates = []
        if preferred_email:
            candidates.append(("per-user", preferred_email, token_storage.get_token("outlook", user_id=preferred_email)))
        candidates.append(("shared", None, token_storage.get_token("outlook")))

        for label, uid, tok in candidates:
            if not tok or not tok.get("access_token"):
                continue

            access_token = tok["access_token"]
            refresh_token = tok.get("refresh_token")

            ok, user_data, status, extra = _graph_me(access_token)
            if ok:
                session.permanent = True
                session["access_token"] = access_token
                session["user_email"] = (user_data.get("mail") or user_data.get("userPrincipalName") or preferred_email)
                session["user_name"] = user_data.get("displayName") or ""
                session.modified = True
                print(f"[AUTH] Auto-restored Outlook ({label}) for {session['user_email']}")
                return True, session["user_email"], session["user_name"]

            # Token invalid/expired - try refresh if we have refresh token
            if refresh_token and status == 401:
                print(f"[AUTH] Outlook token expired ({label}), attempting refresh...")
                try:
                    refresh_result = outlook_refresh_token(refresh_token)
                    
                    new_access = refresh_result.get("access_token")
                    new_refresh = refresh_result.get("refresh_token", refresh_token)
                    
                    if new_access:
                        # Validate the new token
                        ok, user_data, status, extra = _graph_me(new_access)
                        if ok:
                            session.permanent = True
                            session["access_token"] = new_access
                            session["user_email"] = (user_data.get("mail") or user_data.get("userPrincipalName") or preferred_email)
                            session["user_name"] = user_data.get("displayName") or ""
                            session.modified = True
                            
                            # Save refreshed token to Firestore
                            user_email = session["user_email"]
                            token_storage.save_token(
                                service="outlook",
                                access_token=new_access,
                                refresh_token=new_refresh,
                                updated_by=user_email,
                                user_id=user_email
                            )
                            print(f"[AUTH] Auto-refreshed Outlook ({label}) for {user_email}")
                            return True, user_email, session["user_name"]
                        else:
                            print(f"[AUTH] Refreshed Outlook token still invalid: status={status}")
                    else:
                        print(f"[AUTH] Outlook refresh failed: {refresh_result.get('error_description', 'Unknown error')}")
                except Exception as e:
                    print(f"[AUTH] Outlook refresh error: {e}")

            print(f"[AUTH] Stored Outlook token invalid ({label}) status={status} extra={extra}")

        return False, None, None

    def _restore_zoho_from_firestore(user_email):
        """
        Restore Zoho from Firestore - PER-USER ONLY, no shared fallback.
        Returns connected(bool).
        """
        # Already in session
        if session.get("zoho_access_token"):
            return True

        # Per-user token ONLY - no shared fallback (that causes permission issues)
        if not user_email:
            print("[AUTH] No user email - cannot restore Zoho")
            return False
            
        stored = token_storage.get_token("zoho", user_id=user_email)
        if not stored:
            print(f"[AUTH] No Zoho token found for {user_email} - user needs to connect Zoho")
            return False

        access = stored.get("access_token")
        refresh = stored.get("refresh_token")

        # If we have an access token, validate quickly
        if access:
            try:
                test = zoho.get_portals(access)
                if test.get("portals") or test.get("code") == 0:
                    session["zoho_access_token"] = access
                    session["zoho_refresh_token"] = refresh
                    session.modified = True
                    print(f"[AUTH] Auto-restored Zoho for {user_email}")
                    return True
                else:
                    print(f"[AUTH] Zoho access token invalid, will try refresh")
            except Exception as e:
                print(f"[AUTH] Zoho validate error, will try refresh: {e}")

        # If access is bad/missing but refresh exists, refresh and validate
        if refresh:
            try:
                zoho_result = zoho.refresh_access_token(refresh)
                new_access = zoho_result.get("access_token")
                new_refresh = zoho_result.get("refresh_token", refresh)

                if new_access:
                    test = zoho.get_portals(new_access)
                    if test.get("portals") or test.get("code") == 0:
                        session["zoho_access_token"] = new_access
                        session["zoho_refresh_token"] = new_refresh
                        session.modified = True

                        token_storage.save_token(
                            service="zoho",
                            access_token=new_access,
                            refresh_token=new_refresh,
                            updated_by=user_email,
                            user_id=user_email
                        )
                        print(f"[AUTH] Auto-refreshed Zoho for {user_email}")
                        return True

                    print(f"[AUTH] Zoho refresh produced invalid token")
            except Exception as e:
                print(f"[AUTH] Zoho refresh failed: {e}")

        print(f"[AUTH] Zoho restore failed for {user_email} - user needs to reconnect")
        return False

    def _check_shared_mailbox():
        """
        Check if shared mailbox (tom@procalcs.net) token is valid.
        Returns (connected: bool, can_refresh: bool)
        """
        try:
            from outlook_integration import SHARED_MAILBOX, refresh_access_token
        except ImportError:
            from .outlook_integration import SHARED_MAILBOX, refresh_access_token
        
        # Try to get shared mailbox token
        stored = token_storage.get_shared_mailbox_token(SHARED_MAILBOX, "outlook")
        if not stored:
            # Fallback: try old location (tom's user token)
            stored = token_storage.get_token("outlook", user_id=SHARED_MAILBOX)
        
        if not stored:
            print(f"[AUTH] No shared mailbox token found for {SHARED_MAILBOX}")
            return False
        
        access_token = stored.get("access_token")
        refresh_token = stored.get("refresh_token")
        
        if access_token:
            # Validate token
            ok, user_data, status, extra = _graph_me(access_token)
            if ok:
                print(f"[AUTH] Shared mailbox token valid for {SHARED_MAILBOX}")
                return True
            
            # Try refresh if expired
            if refresh_token and status == 401:
                print(f"[AUTH] Shared mailbox token expired, refreshing...")
                try:
                    refresh_result = refresh_access_token(refresh_token)
                    new_access = refresh_result.get("access_token")
                    new_refresh = refresh_result.get("refresh_token", refresh_token)
                    
                    if new_access:
                        ok, user_data, status, extra = _graph_me(new_access)
                        if ok:
                            # Save refreshed token
                            token_storage.save_shared_mailbox_token(
                                mailbox_email=SHARED_MAILBOX,
                                service="outlook",
                                access_token=new_access,
                                refresh_token=new_refresh,
                                updated_by="system"
                            )
                            print(f"[AUTH] Refreshed shared mailbox token for {SHARED_MAILBOX}")
                            return True
                except Exception as e:
                    print(f"[AUTH] Shared mailbox refresh error: {e}")
        
        return False

    # ---------- MAIN FLOW ----------
    # Priority: 1) URL param, 2) Session, 3) Cookie
    preferred_email = (
        request.args.get("email") or 
        session.get("user_email") or 
        request.cookies.get("procalcs_uid")
    )
    
    print(f"[AUTH] Status check - preferred_email: {preferred_email}, has session token: {bool(session.get('access_token'))}, has cookie: {bool(request.cookies.get('procalcs_uid'))}, cookie_value: {request.cookies.get('procalcs_uid')}")

    email = None
    name = None
    outlook_valid = False

    # 1) If session token exists, validate it (don't assume it's good)
    if session.get("access_token"):
        ok, user_data, status, extra = _graph_me(session["access_token"])
        if ok:
            outlook_valid = True
            session["user_email"] = session.get("user_email") or (user_data.get("mail") or user_data.get("userPrincipalName"))
            session["user_name"] = session.get("user_name") or (user_data.get("displayName") or "")
            session.modified = True
            email = session["user_email"]
            name = session["user_name"]
        else:
            print(f"[AUTH] Session Outlook token invalid status={status} extra={extra}; clearing session token")
            session.pop("access_token", None)

    # 2) Restore from Firestore if needed
    if not outlook_valid:
        outlook_valid, email, name = _restore_outlook_from_firestore(preferred_email)

    # 3) Check shared mailbox status (for Front Office)
    shared_mailbox_connected = _check_shared_mailbox()

    # 4) If user Outlook valid, restore Zoho
    zoho_connected = False
    if outlook_valid:
        zoho_connected = _restore_zoho_from_firestore(email)

        # Cache validation timestamp and shared mailbox status for faster subsequent calls
        session["_auth_validated_at"] = time.time()
        session["_shared_mailbox_valid"] = shared_mailbox_connected
        session.modified = True

        permissions = user_roles.get_user_permissions(email)
        return jsonify({
            "authenticated": True,
            "email": email,
            "name": name,
            "role": permissions["role"],
            "allowed_tabs": permissions["allowed_tabs"],
            "zoho_connected": zoho_connected,
            "shared_mailbox_connected": shared_mailbox_connected
        })

    # Even if user not authenticated, return shared mailbox status
    return jsonify({
        "authenticated": False,
        "shared_mailbox_connected": shared_mailbox_connected
    })

@app.route('/auth/logout')
def auth_logout():
    """Complete logout - clears session, cookies, and optionally Firestore tokens"""
    user_email = session.get('user_email')
    cookie_email = request.cookies.get('procalcs_uid')
    
    # Use whichever email we can find
    email_to_clear = user_email or cookie_email
    
    print(f"[AUTH] Logout requested for {email_to_clear}")
    
    # Check if user wants to keep tokens (for quick re-login) or full clear
    full_clear = request.args.get('full', 'false').lower() == 'true'
    
    if full_clear and email_to_clear:
        # Full logout - delete Firestore tokens too
        token_storage.delete_token('outlook', user_id=email_to_clear)
        token_storage.delete_token('zoho', user_id=email_to_clear)
        print(f"[AUTH] Deleted Firestore tokens for {email_to_clear}")
    
    # Always clear session
    session.clear()
    
    # Always clear cookie
    response = make_response(jsonify({'success': True, 'cleared_user': email_to_clear}))
    response.delete_cookie('procalcs_uid', path='/')
    
    print(f"[AUTH] Logout complete for {email_to_clear}")
    return response

# ==================== EMAIL ROUTES ====================
# Routes moved to email_routes.py


# ==================== TEAM LEADER TASKS ====================
# Routes moved to team_tasks_routes.py

@app.route('/api/projects/save-review', methods=['POST'])
def save_review_to_workdrive():
    """Save AI review as Word document to Working Drawings folder"""
    token = get_zoho_token()
    if not token:
        return jsonify({'error': 'Not connected to Zoho'}), 401
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        from datetime import datetime
        
        data = request.json
        project_name = data.get('project_name')
        review_data = data.get('review_data', {})
        designer_name = data.get('designer_name', 'Unknown')
        
        if not project_name:
            return jsonify({'error': 'project_name required'}), 400
        
        access_token = token
        
        # Get client name from Zoho Projects (same as email save)
        portal_id = '655665654'
        projects_response = zoho.get_projects(access_token, portal_id)
        projects = projects_response.get('projects', [])
        project = next((p for p in projects if p.get('name') == project_name), None)
        
        if not project:
            return jsonify({'error': f'Project "{project_name}" not found'}), 404
        
        client_name = None
        custom_fields = project.get('custom_fields', [])
        for field in custom_fields:
            if 'CRM Account Name' in field:
                client_name = field['CRM Account Name']
                break
        
        if not client_name:
            return jsonify({'error': 'No CRM Account Name found for project'}), 404
        
        # Find Working Drawings folder (same pattern as find_emails_folder)
        wd_folder_id = zoho.find_working_drawings_folder(access_token, client_name, project_name)
        if not wd_folder_id:
            return jsonify({'error': 'Working Drawings folder not found'}), 404
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading('AI Project Review', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Project info
        doc.add_paragraph(f"Project: {project_name}")
        doc.add_paragraph(f"Client: {client_name}")
        doc.add_paragraph(f"Reviewed by: {designer_name}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph("")
        
        # Summary section
        summary = review_data.get('summary', {})
        if summary:
            doc.add_heading('Project Summary', level=1)
            for key, value in summary.items():
                if value and value != 'N/A' and value != 'Not found':
                    label = key.replace('_', ' ').title()
                    doc.add_paragraph(f"{label}: {value}")
        
        # Missing info section
        missing = review_data.get('missing_info', [])
        if missing:
            doc.add_heading('Missing Information', level=1)
            for item in missing:
                field = item.get('field', item.get('item', 'Unknown'))
                reason = item.get('reason', '')
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{field}").bold = True
                if reason:
                    p.add_run(f" - {reason}")
        
        # Designer notes section
        notes = review_data.get('designer_notes', [])
        if notes:
            doc.add_heading('Designer Notes', level=1)
            for note in notes:
                doc.add_paragraph(note, style='List Bullet')
        
        # Conflicts section
        conflicts = review_data.get('conflicts', [])
        if conflicts:
            doc.add_heading('Conflicts Found', level=1)
            for conflict in conflicts:
                field = conflict.get('field', 'Unknown')
                values = conflict.get('values', [])
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{field}: ").bold = True
                p.add_run(', '.join(values))
        
        # Files analyzed section
        files = review_data.get('files_analyzed', [])
        if files:
            doc.add_heading('Files Analyzed', level=1)
            for f in files:
                doc.add_paragraph(f, style='List Bullet')
        
        # Save to bytes
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        doc_content = doc_buffer.getvalue()
        
        # Upload to WorkDrive
        timestamp = datetime.now().strftime('%Y-%m-%d')
        filename = f"AI_Review_{timestamp}.docx"
        
        # Use upload function (need to modify for binary)
        import requests as req
        headers = {"Authorization": f"Zoho-oauthtoken {access_token}"}
        files_payload = {'content': (filename, doc_buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        data_payload = {'parent_id': wd_folder_id, 'override-name-exist': 'true'}
        
        upload_url = f"{zoho.ZOHO_WORKDRIVE_API}/upload"
        response = req.post(upload_url, headers=headers, files=files_payload, data=data_payload)
        
        if response.status_code in [200, 201]:
            result = response.json()
            file_id = result.get("data", [{}])[0].get("attributes", {}).get("resource_id")
            file_url = f"https://workdrive.zoho.com/file/{file_id}"
            print(f"[REVIEW] Saved review document: {file_id}")
            return jsonify({
                'success': True,
                'file_id': file_id,
                'file_url': file_url,
                'filename': filename
            })
        else:
            print(f"[REVIEW] Upload failed: {response.status_code} - {response.text[:300]}")
            return jsonify({'error': f'Upload failed: {response.status_code}'}), 500
        
    except ImportError:
        return jsonify({'error': 'python-docx not installed. Run: pip install python-docx'}), 500
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/test-download/<file_id>')
def test_download_endpoints(file_id):
    """Test different download endpoints to find one that works"""
    token = get_zoho_token()
    if not token:
        return jsonify({'error': 'Not connected to Zoho'}), 401
    
    import requests
    access_token = token
    headers = {"Authorization": f"Zoho-oauthtoken {access_token}"}
    
    endpoints = [
        f"https://www.zohoapis.com/workdrive/api/v1/files/{file_id}/download",
        f"https://www.zohoapis.com/workdrive/api/v1/download/{file_id}",
        f"https://download.zoho.com/v1/workdrive/download/{file_id}",
        f"https://workdrive.zoho.com/api/v1/download/{file_id}",
        f"https://workdrive.zoho.com/api/v1/files/{file_id}/download",
    ]
    
    results = []
    for url in endpoints:
        try:
            resp = requests.get(url, headers=headers, timeout=10)
            results.append({
                'url': url,
                'status': resp.status_code,
                'success': resp.status_code == 200,
                'content_length': len(resp.content) if resp.status_code == 200 else 0,
                'response': resp.text[:150] if resp.status_code != 200 else 'OK'
            })
        except Exception as e:
            results.append({'url': url, 'status': 'error', 'response': str(e)})
    
    return jsonify({'file_id': file_id, 'results': results})


@app.route('/api/workdrive/save-email', methods=['POST'])
def save_email_to_workdrive():
    """Save email content to WorkDrive Emails folder"""
    token = get_zoho_token()
    if not token:
        return jsonify({'error': 'Not connected to Zoho'}), 401
    
    data = request.json
    project_name = data.get('project_name')
    filename = data.get('filename')
    content = data.get('content')
    
    if not project_name or not filename or not content:
        return jsonify({'error': 'project_name, filename, and content required'}), 400
    
    access_token = token
    
    # Get client name from Zoho Projects
    portal_id = '655665654'
    projects_response = zoho.get_projects(access_token, portal_id)
    projects = projects_response.get('projects', [])
    project = next((p for p in projects if p.get('name') == project_name), None)
    
    if not project:
        return jsonify({'error': f'Project "{project_name}" not found'}), 404
    
    client_name = None
    custom_fields = project.get('custom_fields', [])
    for field in custom_fields:
        if 'CRM Account Name' in field:
            client_name = field['CRM Account Name']
            break
    
    if not client_name:
        return jsonify({'error': 'No CRM Account Name found for project'}), 404
    
    # Find Emails folder in WorkDrive
    try:
        emails_folder_id = zoho.find_emails_folder(access_token, client_name, project_name)
        if not emails_folder_id:
            return jsonify({'error': 'Could not find Emails folder in WorkDrive'}), 404
        
        # Upload file
        result = zoho.upload_file_to_workdrive(access_token, emails_folder_id, filename, content)
        
        if result:
            return jsonify({'success': True, 'file_id': result})
        else:
            return jsonify({'success': False, 'error': 'Failed to upload file'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ==================== DWG TO PDF CONVERTER ====================

@app.route('/api/workdrive/scan-dwg/<project_id>')
def scan_for_dwg_files(project_id):
    """Scan a project folder recursively for DWG files"""
    token = get_zoho_token()
    if not token:
        return jsonify({'error': 'Not connected to Zoho'}), 401
    
    try:
        dwg_files = []
        
        def scan_folder(folder_id, folder_name='Root'):
            """Recursively scan folder for DWG files"""
            contents = zoho.get_folder_contents(token, folder_id, limit=200)
            for item in contents.get('data', []):
                attrs = item.get('attributes', {})
                name = attrs.get('name', '')
                
                if attrs.get('is_folder'):
                    # Recurse into subfolders
                    scan_folder(item['id'], name)
                elif name.lower().endswith('.dwg'):
                    dwg_files.append({
                        'id': item['id'],
                        'name': name,
                        'folder': folder_name,
                        'folder_id': folder_id
                    })
        
        scan_folder(project_id)
        
        return jsonify({
            'dwg_files': dwg_files,
            'count': len(dwg_files)
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/workdrive/convert-dwg', methods=['POST'])
def convert_dwg_to_pdf():
    """Convert a DWG file to PDF using ODA File Converter"""
    token = get_zoho_token()
    if not token:
        return jsonify({'error': 'Not connected to Zoho'}), 401
    
    try:
        data = request.json
        file_id = data.get('file_id')
        file_name = data.get('file_name')
        folder_id = data.get('folder_id')
        
        if not file_id or not file_name:
            return jsonify({'error': 'file_id and file_name required'}), 400
        
        # Download the DWG file
        print(f"[DWG] Downloading {file_name}...")
        dwg_content = zoho.download_file_content(token, file_id)
        if not dwg_content:
            return jsonify({'error': 'Failed to download DWG file'}), 500
        
        import tempfile
        import subprocess
        import os
        
        # Create temp directory for conversion
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = os.path.join(temp_dir, 'input')
            output_path = os.path.join(temp_dir, 'output')
            os.makedirs(input_path)
            os.makedirs(output_path)
            
            # Write DWG file
            dwg_path = os.path.join(input_path, file_name)
            with open(dwg_path, 'wb') as f:
                f.write(dwg_content)
            
            # Convert using ODA File Converter
            # ODA command: ODAFileConverter "input_folder" "output_folder" ACAD2018 DWG 0 1 "*.dwg"
            # For PDF output, we use different parameters
            
            # Check if ODA is installed
            oda_paths = [
                '/usr/bin/ODAFileConverter',  # Linux
                '/opt/ODAFileConverter/ODAFileConverter',  # Linux alternative
                'C:\\Program Files\\ODA\\ODAFileConverter\\ODAFileConverter.exe',  # Windows
                '/Applications/ODAFileConverter.app/Contents/MacOS/ODAFileConverter'  # Mac
            ]
            
            oda_path = None
            for p in oda_paths:
                if os.path.exists(p):
                    oda_path = p
                    break
            
            if not oda_path:
                # ODA not installed - provide helpful message
                # TODO: Consider CloudConvert API or similar service for production
                return jsonify({
                    'error': 'DWG to PDF conversion is not yet available on the server. This feature requires additional setup.',
                    'suggestion': 'For now, please convert DWG files to PDF using AutoCAD or a free online converter like zamzar.com'
                }), 501
            
            # Run ODA converter
            # Arguments: InputFolder OutputFolder OutputVersion OutputType RecurseOption AuditOption [FilterFile]
            # OutputType: DWG=0, DXF=1, DXB=2, PDF=3
            result = subprocess.run([
                oda_path,
                input_path,
                output_path,
                'ACAD2018',
                'PDF',
                '0',  # No recursion
                '1',  # Audit
            ], capture_output=True, text=True, timeout=120)
            
            print(f"[DWG] ODA output: {result.stdout}")
            if result.returncode != 0:
                print(f"[DWG] ODA error: {result.stderr}")
            
            # Find the output PDF
            pdf_name = file_name.rsplit('.', 1)[0] + '.pdf'
            pdf_path = os.path.join(output_path, pdf_name)
            
            if not os.path.exists(pdf_path):
                # Check for any PDF in output
                for f in os.listdir(output_path):
                    if f.lower().endswith('.pdf'):
                        pdf_path = os.path.join(output_path, f)
                        pdf_name = f
                        break
                else:
                    return jsonify({
                        'error': 'Conversion failed - no PDF generated',
                        'oda_output': result.stdout,
                        'oda_error': result.stderr
                    }), 500
            
            # Read the PDF
            with open(pdf_path, 'rb') as f:
                pdf_content = f.read()
            
            # Upload PDF to same folder in WorkDrive
            print(f"[DWG] Uploading {pdf_name} to WorkDrive...")
            upload_result = zoho.upload_binary_to_workdrive(token, folder_id, pdf_name, pdf_content)
            
            if not upload_result:
                return jsonify({'error': 'Failed to upload PDF to WorkDrive'}), 500
            
            return jsonify({
                'success': True,
                'pdf_name': pdf_name,
                'pdf_id': upload_result.get('id'),
                'message': f'Successfully converted {file_name} to {pdf_name}'
            })
            
    except subprocess.TimeoutExpired:
        return jsonify({'error': 'Conversion timed out after 120 seconds'}), 500
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/outlook/convert-dwg-attachment', methods=['POST'])
def convert_dwg_email_attachment():
    """Download DWG attachment from email and convert to PDF"""
    access_token = get_outlook_token()
    if not access_token:
        return jsonify({'error': 'Not connected to Outlook'}), 401
    
    try:
        data = request.json
        email_id = data.get('email_id')
        attachment_id = data.get('attachment_id')
        attachment_name = data.get('attachment_name')
        
        if not email_id or not attachment_id or not attachment_name:
            return jsonify({'error': 'email_id, attachment_id, and attachment_name required'}), 400
        
        # Download the attachment from Outlook
        print(f"[DWG] Downloading attachment {attachment_name} from email...")
        
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json"
        }
        
        # Get attachment content
        attachment_url = f"https://graph.microsoft.com/v1.0/me/messages/{email_id}/attachments/{attachment_id}"
        response = requests.get(attachment_url, headers=headers)
        
        if response.status_code != 200:
            return jsonify({'error': f'Failed to download attachment: {response.status_code}'}), 500
        
        attachment_data = response.json()
        content_bytes = attachment_data.get('contentBytes')
        
        if not content_bytes:
            return jsonify({'error': 'Attachment has no content'}), 500
        
        import base64
        dwg_content = base64.b64decode(content_bytes)
        
        # For now, return a message since ODA converter is not installed
        # TODO: Implement actual conversion when ODA is available
        return jsonify({
            'error': 'DWG to PDF conversion is not yet available. Please use AutoCAD or a free online converter like zamzar.com to convert this file.',
            'suggestion': f'Download {attachment_name} and convert it manually, then upload the PDF to the project folder.'
        }), 501
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

# ==================== END DWG TO PDF CONVERTER ====================


@app.route('/api/workdrive/all-projects')
def get_all_projects():
    """Search projects by name - searches client names AND project names"""
    token = get_zoho_token()
    if not token:
        return jsonify({'error': 'Not connected to Zoho'}), 401
    
    query = request.args.get('q', '').lower().strip()
    if not query or len(query) < 2:
        return jsonify({'projects': [], 'message': 'Type at least 2 characters to search'})
    
    try:
        # Get all clients from cache (fast)
        cached_clients = get_cached_clients(token)
        
        # Filter to clients whose name contains the query
        matching_clients = [c for c in cached_clients if query in c['attributes']['name'].lower()]
        
        # Limit to first 30 matching clients to avoid too many API calls
        matching_clients = matching_clients[:30]
        
        print(f"[ProjectSearch] Query '{query}' matched {len(matching_clients)} clients by name")
        
        matching_projects = []
        searched_client_ids = set()
        
        # Search projects in clients that match by name
        for client in matching_clients:
            client_id = client['id']
            searched_client_ids.add(client_id)
            client_name = client['attributes']['name']
            
            folders = zoho.get_folder_contents(token, client_id, limit=100)
            for f in folders.get('data', []):
                if f.get('attributes', {}).get('is_folder'):
                    project_name = f['attributes']['name']
                    matching_projects.append({
                        'id': f['id'],
                        'name': project_name,
                        'client_id': client_id,
                        'client_name': client_name,
                        'display_name': f"{project_name} ({client_name})"
                    })
        
        # If we found projects from matching client names, return them
        # Otherwise search more clients for project names
        if len(matching_projects) < 5:
            # Search more clients for projects whose NAME matches the query
            # Limit to checking 50 clients to balance speed vs coverage
            clients_to_search = [c for c in cached_clients if c['id'] not in searched_client_ids][:50]
            
            for client in clients_to_search:
                client_id = client['id']
                client_name = client['attributes']['name']
                
                folders = zoho.get_folder_contents(token, client_id, limit=100)
                for f in folders.get('data', []):
                    if f.get('attributes', {}).get('is_folder'):
                        project_name = f['attributes']['name']
                        # Check if query matches project name
                        if query in project_name.lower():
                            matching_projects.append({
                                'id': f['id'],
                                'name': project_name,
                                'client_id': client_id,
                                'client_name': client_name,
                                'display_name': f"{project_name} ({client_name})"
                            })
                
                # Stop if we have enough results
                if len(matching_projects) >= 20:
                    break
        
        # Remove duplicates and sort
        seen = set()
        unique_projects = []
        for p in matching_projects:
            if p['id'] not in seen:
                seen.add(p['id'])
                unique_projects.append(p)
        
        print(f"[ProjectSearch] Found {len(unique_projects)} total projects")
        unique_projects.sort(key=lambda x: x['name'].lower())
        return jsonify({'projects': unique_projects[:50], 'count': len(unique_projects)})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/frontoffice/file-email', methods=['POST'])
def frontoffice_file_email():
    """File email and attachments to a project folder or subfolder"""
    zoho_token = get_zoho_token()
    if not zoho_token:
        return jsonify({'error': 'Not connected to Zoho'}), 401
    outlook_token = get_outlook_token()
    if not outlook_token:
        return jsonify({'error': 'Not connected to Outlook'}), 401
    
    data = request.json
    email_id = data.get('email_id')
    project_folder_id = data.get('project_folder_id')
    subfolder_name = data.get('subfolder_name')  # Will be set if user selected a subfolder
    
    if not email_id or not project_folder_id:
        return jsonify({'error': 'email_id and project_folder_id required'}), 400
    
    zoho_token = session['zoho_access_token']
    outlook_token = session['access_token']
    
    try:
        # Determine target folder:
        # - If subfolder_name is provided, user selected a specific subfolder - use project_folder_id directly (it's the subfolder)
        # - If no subfolder_name, user selected project root - find Emails folder inside it
        if subfolder_name:
            # User selected a specific subfolder - file directly to it
            target_folder_id = project_folder_id
            print(f"[FileEmail] Filing to selected subfolder: {subfolder_name}")
        else:
            # User selected project root - find Emails folder
            emails_folder = zoho.find_subfolder(zoho_token, project_folder_id, 'Emails')
            if not emails_folder:
                # Try 'EMails' (different capitalization)
                emails_folder = zoho.find_subfolder(zoho_token, project_folder_id, 'EMails')
            if not emails_folder:
                return jsonify({'error': 'Emails folder not found in project'}), 404
            target_folder_id = emails_folder.get('id')
            print(f"[FileEmail] Filing to auto-found Emails folder")
        
        # Get full email from Outlook
        email = get_email_details(outlook_token, email_id)
        
        if email.get('error'):
            return jsonify({'error': f"Failed to get email: {email.get('error', {}).get('message', 'Unknown')}"}), 500
        
        # Save email as text file
        from datetime import datetime
        timestamp = datetime.now().strftime('%Y-%m-%d_%H%M%S')
        subject = email.get('subject', 'No Subject')[:50].replace('/', '-').replace('\\', '-')
        email_filename = f"Email_{timestamp}_{subject}.txt"
        
        # Build email content
        from_addr = email.get('from', {}).get('emailAddress', {})
        email_content = f"""From: {from_addr.get('name', '')} <{from_addr.get('address', '')}>
Subject: {email.get('subject', '')}
Date: {email.get('receivedDateTime', '')}

{email.get('body', {}).get('content', email.get('bodyPreview', ''))}
"""
        # Strip HTML if present
        if email.get('body', {}).get('contentType') == 'html':
            import re
            email_content = re.sub('<[^<]+?>', '', email_content)
        
        # Upload email text
        zoho.upload_file_to_workdrive(zoho_token, target_folder_id, email_filename, email_content)
        
        # Upload attachments
        attachments = email.get('attachments', [])
        uploaded_attachments = []
        
        for att in attachments:
            if att.get('contentBytes'):
                import base64
                att_content = base64.b64decode(att['contentBytes'])
                att_name = att.get('name', 'attachment')
                
                # Upload binary attachment
                result = zoho.upload_binary_to_workdrive(zoho_token, target_folder_id, att_name, att_content)
                if result:
                    uploaded_attachments.append(att_name)
        
        # Track email filed
        client_name = data.get('client_name', 'Unknown')
        project_name = data.get('project_name', 'Unknown')
        track_email_filed(email_id, project_name, client_name)
        
        # Move email from Inbox to "Filed" folder in Outlook
        try:
            from outlook_integration import SHARED_MAILBOX
        except ImportError:
            from .outlook_integration import SHARED_MAILBOX
        
        headers = {'Authorization': f'Bearer {outlook_token}'}
        
        # Get or create "Filed" folder
        folders_url = f'https://graph.microsoft.com/v1.0/users/{SHARED_MAILBOX}/mailFolders'
        folders_resp = requests.get(folders_url, headers=headers)
        filed_folder_id = None
        
        if folders_resp.status_code == 200:
            for folder in folders_resp.json().get('value', []):
                if folder.get('displayName', '').lower() == 'filed':
                    filed_folder_id = folder.get('id')
                    break
        
        # Create "Filed" folder if it doesn't exist
        if not filed_folder_id:
            create_resp = requests.post(folders_url, headers=headers, json={'displayName': 'Filed'})
            if create_resp.status_code == 201:
                filed_folder_id = create_resp.json().get('id')
                print(f"[FileEmail] Created 'Filed' folder in Outlook")
        
        # Move email to Filed folder
        email_moved = False
        if filed_folder_id:
            move_url = f'https://graph.microsoft.com/v1.0/users/{SHARED_MAILBOX}/messages/{email_id}/move'
            move_resp = requests.post(move_url, headers=headers, json={'destinationId': filed_folder_id})
            if move_resp.status_code == 201:
                email_moved = True
                print(f"[FileEmail] Moved email to 'Filed' folder")
            else:
                print(f"[FileEmail] Failed to move email: {move_resp.status_code} - {move_resp.text}")
        
        return jsonify({
            'success': True,
            'email_saved': email_filename,
            'attachments_saved': uploaded_attachments,
            'email_moved_to_filed': email_moved
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ==================== EMAIL FLAGS (Estimate Requests) ====================

@app.route('/api/frontoffice/email-flags', methods=['GET'])
def get_email_flags():
    """Get all email flags for filtering"""
    try:
        flags_ref = db.collection('email_flags')
        docs = flags_ref.stream()
        flags = {}
        for doc in docs:
            flags[doc.id] = doc.to_dict()
        return jsonify({'success': True, 'flags': flags})
    except Exception as e:
        print(f"[EmailFlags] Error getting flags: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/frontoffice/email-flags', methods=['POST'])
def set_email_flag():
    """Set or update an email flag (estimate, general, etc.)"""
    try:
        data = request.json
        email_id = data.get('email_id')
        flag_type = data.get('flag_type')  # 'estimate', 'general', or None to remove
        
        if not email_id:
            return jsonify({'error': 'email_id required'}), 400
        
        flags_ref = db.collection('email_flags').document(email_id)
        
        if flag_type:
            flags_ref.set({
                'flag_type': flag_type,
                'flagged_by': session.get('user_email', 'unknown'),
                'flagged_at': datetime.now().isoformat()
            })
            print(f"[EmailFlags] Set {email_id[:20]}... as '{flag_type}'")
        else:
            # Remove flag
            flags_ref.delete()
            print(f"[EmailFlags] Removed flag from {email_id[:20]}...")
        
        return jsonify({'success': True})
    except Exception as e:
        print(f"[EmailFlags] Error setting flag: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/frontoffice/create-project', methods=['POST'])
def frontoffice_create_project():
    """Create new project folder structure in WorkDrive"""
    if 'zoho_access_token' not in session:
        return jsonify({'error': 'Not connected to Zoho'}), 401
    
    data = request.json
    client_id = data.get('client_id')  # None if new client
    client_name = data.get('client_name')
    project_name = data.get('project_name')
    email_id = data.get('email_id')  # Optional - file email after creation
    
    if not client_name or not project_name:
        return jsonify({'error': 'client_name and project_name required'}), 400
    
    zoho_token = session['zoho_access_token']
    
    try:
        # If no client_id, create new client folder
        if not client_id:
            print(f"[FrontOffice] Creating new client folder: {client_name}")
            client_id = zoho.create_folder(zoho_token, CLIENTS_FOLDER_ID, client_name)
            if not client_id:
                return jsonify({'error': 'Failed to create client folder'}), 500
            print(f"[FrontOffice] Created client folder: {client_id}")
        
        # Create project folder under client
        print(f"[FrontOffice] Creating project folder: {project_name}")
        project_id = zoho.create_folder(zoho_token, client_id, project_name)
        if not project_id:
            return jsonify({'error': 'Failed to create project folder'}), 500
        print(f"[FrontOffice] Created project folder: {project_id}")
        
        # Create subfolders
        subfolders = ['Files From Client', 'Forms', 'Working Drawings', 'Finals', 'Emails']
        created_folders = {}
        
        for folder_name in subfolders:
            folder_id = zoho.create_folder(zoho_token, project_id, folder_name)
            created_folders[folder_name] = folder_id
            print(f"[FrontOffice] Created subfolder {folder_name}: {folder_id}")
        
        # File email if requested
        email_result = None
        if email_id and 'access_token' in session:
            emails_folder_id = created_folders.get('Emails')
            if emails_folder_id:
                outlook_token = session['access_token']
                
                # Get email details
                email = get_email_details(outlook_token, email_id)
                
                if not email.get('error'):
                    # Save email as text file
                    from datetime import datetime
                    timestamp = datetime.now().strftime('%Y-%m-%d_%H%M%S')
                    subject = email.get('subject', 'No Subject')[:50].replace('/', '-').replace('\\', '-')
                    email_filename = f"Email_{timestamp}_{subject}.txt"
                    
                    from_addr = email.get('from', {}).get('emailAddress', {})
                    email_content = f"""From: {from_addr.get('name', '')} <{from_addr.get('address', '')}>
Subject: {email.get('subject', '')}
Date: {email.get('receivedDateTime', '')}

{email.get('body', {}).get('content', email.get('bodyPreview', ''))}
"""
                    if email.get('body', {}).get('contentType') == 'html':
                        import re
                        email_content = re.sub('<[^<]+?>', '', email_content)
                    
                    zoho.upload_file_to_workdrive(zoho_token, emails_folder_id, email_filename, email_content)
                    
                    # Upload attachments to Files From Client
                    files_folder_id = created_folders.get('Files From Client')
                    attachments = email.get('attachments', [])
                    uploaded = []
                    
                    for att in attachments:
                        if att.get('contentBytes'):
                            import base64
                            att_content = base64.b64decode(att['contentBytes'])
                            att_name = att.get('name', 'attachment')
                            zoho.upload_binary_to_workdrive(zoho_token, files_folder_id, att_name, att_content)
                            uploaded.append(att_name)
                    
                    email_result = {'email_saved': email_filename, 'attachments_saved': uploaded}
                    
                    # Track email filed
                    track_email_filed(email_id, project_name, client_name)
        
        return jsonify({
            'success': True,
            'client_id': client_id,
            'project_id': project_id,
            'folders_created': list(created_folders.keys()),
            'email_filed': email_result
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ==================== EMAIL TRACKING (Firestore) ====================

def load_email_tracking():
    """Load email tracking data from Firestore"""
    try:
        tracking = {"emails": {}, "email_to_client": {}}
        docs = db.collection('email_tracking').stream()
        for doc in docs:
            tracking["emails"][doc.id] = doc.to_dict()
        return tracking
    except Exception as e:
        print(f"[EmailTracking] Error loading from Firestore: {e}")
        return {"emails": {}, "email_to_client": {}}

def save_email_tracking_item(email_id, data):
    """Save single email tracking item to Firestore"""
    try:
        db.collection('email_tracking').document(email_id).set(data)
    except Exception as e:
        print(f"[EmailTracking] Error saving to Firestore: {e}")

def track_email_received(email_id, from_email, subject, received_time):
    """Track when an email is received/seen"""
    tracking = load_email_tracking()
    inquiry_id = None
    
    if email_id not in tracking["emails"]:
        # Auto-create pipeline inquiry for new email addresses
        if from_email:
            inquiry, is_new = get_or_create_inquiry(from_email, received_time, subject)
            inquiry_id = inquiry.get("inquiry_id")
            if is_new:
                print(f"[Pipeline] New inquiry created for {from_email}")
            
            # NEW: Log to lead_events system
            try:
                lead = lead_events.get_or_create_lead(from_email, subject)
                event_type, direction, meta = event_detection.detect_event_type(
                    from_email=from_email,
                    to_email='tom@procalcs.net',  # Default for inbound
                    subject=subject,
                    body='',
                    attachments=[]
                )
                lead_events.log_event(
                    lead_id=lead['id'],
                    event_type=event_type,
                    event_ts=received_time or datetime.now().isoformat(),
                    source_system='outlook',
                    source_id=email_id,
                    direction=direction,
                    meta=meta
                )
            except Exception as e:
                print(f"[LeadEvents] Error logging event: {e}")
        
        tracking["emails"][email_id] = {
            "from_email": from_email,
            "subject": subject,
            "received_at": received_time,
            "first_seen_at": datetime.now().isoformat(),
            "filed_at": None,
            "filed_to_project": None,
            "filed_to_client": None,
            "responded_at": None,
            "inquiry_id": inquiry_id
        }
        save_email_tracking_item(email_id, tracking["emails"][email_id])
    
    return tracking["emails"][email_id]

def track_email_filed(email_id, project_name, client_name):
    """Track when an email is filed to a project"""
    tracking = load_email_tracking()
    
    if email_id in tracking["emails"]:
        tracking["emails"][email_id]["filed_at"] = datetime.now().isoformat()
        tracking["emails"][email_id]["filed_to_project"] = project_name
        tracking["emails"][email_id]["filed_to_client"] = client_name
        save_email_tracking_item(email_id, tracking["emails"][email_id])

def track_email_responded(email_id):
    """Track when an email is responded to"""
    tracking = load_email_tracking()
    
    if email_id in tracking["emails"]:
        tracking["emails"][email_id]["responded_at"] = datetime.now().isoformat()
        save_email_tracking_item(email_id, tracking["emails"][email_id])

@app.route('/api/frontoffice/track-email', methods=['POST'])
def track_email_endpoint():
    """Track an email when it appears in the queue"""
    try:
        data = request.json
        email_id = data.get('email_id')
        from_email = data.get('from_email')
        subject = data.get('subject')
        received_time = data.get('received_time')
        
        if not email_id or not from_email:
            return jsonify({'error': 'email_id and from_email required'}), 400
        
        result = track_email_received(email_id, from_email, subject, received_time)
        return jsonify({'success': True, 'tracking': result})
    except Exception as e:
        print(f"[track-email] Error: {e}")
        return jsonify({'error': str(e), 'success': False}), 500

# Pipeline routes moved to pipeline_routes.py

@app.route('/api/frontoffice/metrics', methods=['GET'])
def get_frontoffice_metrics():
    """Get front office metrics - response time only for NEW conversations"""
    tracking = load_email_tracking()
    emails = tracking.get("emails", {})
    
    now = datetime.now()
    
    # Calculate metrics
    pending_count = 0
    over_24h = 0
    over_48h = 0
    filed_today = 0
    total_filed = 0
    
    today_start = now.replace(hour=0, minute=0, second=0, microsecond=0)
    
    for email_id, data in emails.items():
        if not data.get("filed_at"):
            pending_count += 1
            
            # Check age
            first_seen = datetime.fromisoformat(data.get("first_seen_at", now.isoformat()))
            age_hours = (now - first_seen).total_seconds() / 3600
            
            if age_hours > 48:
                over_48h += 1
            elif age_hours > 24:
                over_24h += 1
        else:
            total_filed += 1
            filed_time = datetime.fromisoformat(data["filed_at"])
            if filed_time >= today_start:
                filed_today += 1
    
    # Calculate average response time from Firestore conversation_responses
    # Only counts FIRST inbound email in each conversation thread
    response_times_minutes = []
    try:
        conv_ref = db.collection('conversation_responses')
        docs = conv_ref.stream()
        for doc in docs:
            data = doc.to_dict()
            inbound_time = data.get('first_inbound_at')
            response_time = data.get('first_response_at')
            if inbound_time and response_time:
                try:
                    t1 = datetime.fromisoformat(inbound_time.replace('Z', ''))
                    t2 = datetime.fromisoformat(response_time.replace('Z', ''))
                    minutes = (t2 - t1).total_seconds() / 60
                    if 0 < minutes < 1440:  # Only count if between 0-24 hours (ignore stale data)
                        response_times_minutes.append(minutes)
                except:
                    pass
    except Exception as e:
        print(f"[Metrics] Error loading conversation responses: {e}")
    
    avg_response_minutes = round(sum(response_times_minutes) / len(response_times_minutes), 1) if response_times_minutes else None
    
    # Color coding: green if under 15min, red if over
    response_status = 'good' if avg_response_minutes and avg_response_minutes <= 15 else 'warning'
    
    return jsonify({
        "pending_count": pending_count,
        "over_24h": over_24h,
        "over_48h": over_48h,
        "filed_today": filed_today,
        "total_filed": total_filed,
        "avg_response_minutes": avg_response_minutes,
        "response_status": response_status,
        "target_minutes": 15
    })


# ==================== PIPELINE TRACKING ====================
# Routes and functions moved to pipeline_routes.py


# ==================== USER SIGNATURES ====================
# Routes moved to user_signatures_routes.py


# ==================== LOCAL FILE REVIEW API ====================

@app.route('/api/review-local-files', methods=['POST'])
def review_local_files():
    """Review local files uploaded by designer"""
    import tempfile
    
    if 'files' not in request.files:
        return jsonify({'error': 'No files uploaded'}), 400
    
    project_name = request.form.get('project_name', 'Unknown Project')
    files = request.files.getlist('files')
    
    if not files or len(files) == 0:
        return jsonify({'error': 'No files selected'}), 400
    
    try:
        # Read file contents
        file_contents = []
        for f in files:
            filename = f.filename
            content = f.read()
            
            # Try to decode as text
            try:
                # .rup files are UTF-16 encoded
                if filename.lower().endswith('.rup'):
                    text_content = content.decode('utf-16', errors='replace')
                else:
                    text_content = content.decode('utf-8', errors='replace')
                
                file_contents.append({
                    'filename': filename,
                    'content': text_content[:50000],  # Limit size
                    'type': 'text'
                })
            except:
                file_contents.append({
                    'filename': filename,
                    'content': f'[Binary file - {len(content)} bytes]',
                    'type': 'binary'
                })
        
        # Build analysis prompt
        files_summary = "\n\n".join([
            f"=== {fc['filename']} ===\n{fc['content']}" 
            for fc in file_contents
        ])
        
        prompt = f"""You are reviewing working HVAC design files for ProCalcs. 
        
Project: {project_name}

The designer has uploaded their LOCAL working files for a mid-project check. 
This is NOT the final submission - they want feedback before finishing.

FILES:
{files_summary}

Please analyze these files and provide:

1. **Quick Status** - Overall impression (Good/Needs Attention/Issues Found)

2. **Key Findings** - What you see in the files (equipment selections, calculations, etc.)

3. **Potential Issues** - Any concerns or things that look incorrect

4. **Suggestions** - Things to double-check before final submission

5. **Missing Information** - What appears to be incomplete

Keep your response practical and actionable. This is a working check, not final QC.
Format with clear headers and bullet points for easy scanning."""

        # Call AI
        client = anthropic.Anthropic(api_key=Config.ANTHROPIC_API_KEY)
        
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}]
        )
        
        analysis = response.content[0].text
        
        # Convert markdown to basic HTML
        analysis_html = analysis.replace('\n\n', '</p><p>').replace('\n', '<br>')
        analysis_html = f'<p>{analysis_html}</p>'
        analysis_html = analysis_html.replace('**', '<strong>').replace('**', '</strong>')
        
        return jsonify({
            'success': True,
            'project': project_name,
            'files_analyzed': [f['filename'] for f in file_contents],
            'analysis': analysis_html
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ============ EMAIL SCANNER API ============

@app.route('/api/scanner/status', methods=['GET'])
def scanner_status():
    """Get email scanner status"""
    status = email_scanner.get_scanner_status()
    # Check token: 1) shared mailbox, 2) user token, 3) legacy shared
    token_data = token_storage.get_shared_mailbox_token('tom@procalcs.net', 'outlook')
    if not token_data:
        user_email = session.get('user_email')
        if user_email:
            token_data = token_storage.get_token('outlook', user_id=user_email)
    if not token_data:
        token_data = token_storage.get_token('outlook')  # Fall back to legacy shared
    status['has_token'] = token_data is not None
    status['token_updated_by'] = token_data.get('updated_by') if token_data else None
    return jsonify(status)

@app.route('/api/scanner/start', methods=['POST'])
def scanner_start():
    """Start the email scanner"""
    email_scanner.start_scanner()
    return jsonify({'success': True, 'message': 'Scanner started'})

@app.route('/api/scanner/stop', methods=['POST'])
def scanner_stop():
    """Stop the email scanner"""
    email_scanner.stop_scanner()
    return jsonify({'success': True, 'message': 'Scanner stopped'})


# ==================== INQUIRY ANALYSIS (AI Auto-Response) ====================

try:
    from . import inquiry_analysis
except ImportError:
    import inquiry_analysis

@app.route('/api/inquiry/analyze', methods=['POST'])
def api_analyze_inquiry():
    """Analyze an email inquiry and draft a response"""
    user_email = session.get('user_email')
    if not user_email:
        return jsonify({'error': 'Not authenticated'}), 401
    
    user_role = user_roles.get_user_role(user_email)
    if user_role not in ['admin', 'ceo', 'coo', 'front_office']:
        return jsonify({'error': 'Not authorized'}), 403
    
    try:
        data = request.get_json()
        message_id = data.get('message_id')
        
        if not message_id:
            return jsonify({'error': 'Missing message_id'}), 400
        
        # Get the email details - use outlook_integration which handles token refresh
        token = get_outlook_token()
        if not token:
            return jsonify({'error': 'Outlook not connected'}), 401
        
        # Fetch email with attachments
        graph_url = f"https://graph.microsoft.com/v1.0/me/messages/{message_id}"
        headers = {"Authorization": f"Bearer {token}"}
        
        email_resp = requests.get(graph_url, headers=headers)
        
        # If token expired, try to refresh
        if email_resp.status_code == 401:
            print("[InquiryAnalysis] Token expired, attempting refresh...")
            stored = token_storage.get_token('outlook', user_id=user_email)
            if stored and stored.get('refresh_token'):
                try:
                    new_tokens = outlook_refresh_token(stored['refresh_token'])
                    if new_tokens and new_tokens.get('access_token'):
                        token = new_tokens['access_token']
                        session['access_token'] = token
                        token_storage.save_token(
                            service='outlook',
                            access_token=token,
                            refresh_token=new_tokens.get('refresh_token', stored['refresh_token']),
                            updated_by=user_email,
                            user_id=user_email
                        )
                        headers = {"Authorization": f"Bearer {token}"}
                        email_resp = requests.get(graph_url, headers=headers)
                except Exception as e:
                    print(f"[InquiryAnalysis] Token refresh failed: {e}")
        
        if email_resp.status_code != 200:
            print(f"[InquiryAnalysis] Email fetch failed: {email_resp.status_code} - {email_resp.text[:200]}")
            return jsonify({'error': f'Failed to fetch email: {email_resp.status_code}'}), 500
        
        email_data = email_resp.json()
        subject = email_data.get('subject', '')
        body = email_data.get('body', {}).get('content', '')
        
        # Strip HTML if present
        if '<html' in body.lower() or '<body' in body.lower():
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(body, 'html.parser')
            body = soup.get_text(separator='\n')
        
        # Get attachments if any
        attachments = []
        if email_data.get('hasAttachments'):
            att_url = f"{graph_url}/attachments"
            att_resp = requests.get(att_url, headers=headers)
            if att_resp.status_code == 200:
                for att in att_resp.json().get('value', []):
                    if att.get('@odata.type') == '#microsoft.graph.fileAttachment':
                        attachments.append({
                            'name': att.get('name'),
                            'content_type': att.get('contentType'),
                            'data': att.get('contentBytes')  # Already base64
                        })
        
        print(f"[InquiryAnalysis] Analyzing email: {subject}")
        print(f"[InquiryAnalysis] Attachments: {[a['name'] for a in attachments]}")
        
        # Analyze the inquiry
        analysis = inquiry_analysis.analyze_inquiry(subject, body, attachments)
        
        if 'error' in analysis:
            return jsonify({'error': analysis['error']}), 500
        
        # Draft response
        draft = inquiry_analysis.draft_response(analysis)
        
        if 'error' in draft:
            return jsonify({'error': draft['error']}), 500
        
        # Get sender info for reply
        sender = email_data.get('from', {}).get('emailAddress', {})
        
        return jsonify({
            'success': True,
            'analysis': analysis,
            'draft': {
                'to': sender.get('address'),
                'to_name': sender.get('name'),
                'subject': draft['subject'],
                'body': draft['body'],
                'attachments_to_include': draft['attachments_to_include']
            },
            'original_message_id': message_id
        })
        
    except Exception as e:
        print(f"[InquiryAnalysis] Error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/inquiry/send', methods=['POST'])
def api_send_inquiry_response():
    """Send the drafted response and move original to AI Auto Response folder"""
    user_email = session.get('user_email')
    if not user_email:
        return jsonify({'error': 'Not authenticated'}), 401
    
    user_role = user_roles.get_user_role(user_email)
    if user_role not in ['admin', 'ceo', 'coo', 'front_office']:
        return jsonify({'error': 'Not authorized'}), 403
    
    try:
        data = request.get_json()
        to_email = data.get('to')
        subject = data.get('subject')
        body = data.get('body')
        attachments_to_include = data.get('attachments_to_include', [])
        original_message_id = data.get('original_message_id')
        
        if not all([to_email, subject, body]):
            return jsonify({'error': 'Missing required fields'}), 400
        
        token = get_outlook_token()
        if not token:
            return jsonify({'error': 'Outlook not connected'}), 401
        
        headers = {
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json"
        }
        
        # Build email message
        message = {
            "subject": subject,
            "body": {
                "contentType": "HTML",
                "content": body.replace('\n', '<br>')
            },
            "toRecipients": [{"emailAddress": {"address": to_email}}]
        }
        
        # TODO: Add WorkDrive attachments
        # For now, we'll note which files should be attached
        # Full implementation would download from WorkDrive and attach
        
        if attachments_to_include:
            print(f"[InquiryAnalysis] Should attach: {attachments_to_include}")
            # Add note to email about attachments
            attachment_note = "\n\n---\nAttachments included: "
            file_config = inquiry_analysis.get_attachment_files().get('files', {})
            for att_key in attachments_to_include:
                filename = file_config.get(att_key, att_key)
                attachment_note += f"\n• {filename}"
            message['body']['content'] += attachment_note.replace('\n', '<br>')
        
        # Send the email
        send_url = "https://graph.microsoft.com/v1.0/me/sendMail"
        send_resp = requests.post(send_url, headers=headers, json={"message": message})
        
        if send_resp.status_code not in [200, 202]:
            return jsonify({'error': f'Failed to send: {send_resp.text}'}), 500
        
        print(f"[InquiryAnalysis] Email sent to {to_email}")
        
        # Move original email to "AI Auto Response" folder
        if original_message_id:
            try:
                from outlook_integration import resolve_folder_id
            except ImportError:
                from .outlook_integration import resolve_folder_id
            
            # Use centralized resolver (handles pagination)
            target_folder_id = resolve_folder_id(token, 'AI Auto Response', None)
            
            if target_folder_id:
                move_url = f"https://graph.microsoft.com/v1.0/me/messages/{original_message_id}/move"
                move_resp = requests.post(move_url, headers=headers, json={"destinationId": target_folder_id})
                if move_resp.status_code == 201:
                    print(f"[InquiryAnalysis] Moved to AI Auto Response folder")
                else:
                    print(f"[InquiryAnalysis] Failed to move: {move_resp.text}")
            else:
                print(f"[InquiryAnalysis] AI Auto Response folder not found")
        
        # Track conversation response time
        original_received_at = data.get('original_received_at')
        if original_message_id and original_received_at:
            try:
                conv_ref = db.collection('conversation_responses').document(original_message_id)
                existing = conv_ref.get()
                if not existing.exists:
                    conv_ref.set({
                        'first_inbound_at': original_received_at,
                        'first_response_at': datetime.now().isoformat(),
                        'to_email': to_email,
                        'subject': subject,
                        'responded_by': user_email
                    })
                    print(f"[ResponseTime] Tracked AI response to {original_message_id[:20]}...")
            except Exception as track_err:
                print(f"[ResponseTime] Error tracking: {track_err}")
        
        return jsonify({'success': True, 'message': 'Email sent and original moved'})
        
    except Exception as e:
        print(f"[InquiryAnalysis] Send error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ==================== ESTIMATE ANALYSIS ====================

@app.route('/api/estimate/analyze', methods=['POST'])
def api_analyze_for_estimate():
    """Analyze an email to extract estimate details"""
    user_email = session.get('user_email')
    if not user_email:
        return jsonify({'error': 'Not authenticated'}), 401
    
    try:
        data = request.get_json()
        message_id = data.get('message_id')
        
        if not message_id:
            return jsonify({'error': 'Missing message_id'}), 400
        
        import base64  # Ensure base64 is available for attachment processing
        
        # Get Outlook token
        token = get_outlook_token()
        if not token:
            return jsonify({'error': 'Not connected to Outlook'}), 401
        
        # Check if using shared mailbox (non-designer roles)
        user_role = user_roles.get_user_role(user_email) if user_email else 'designer'
        shared_mailbox = None
        # All roles except designer and payroll need access to shared mailbox
        if user_role in ['front_office', 'admin', 'ceo', 'coo', 'team_leader']:
            try:
                from outlook_integration import SHARED_MAILBOX
            except ImportError:
                from .outlook_integration import SHARED_MAILBOX
            shared_mailbox = SHARED_MAILBOX
        
        headers = {"Authorization": f"Bearer {token}"}
        
        # Get email details - use shared mailbox if applicable
        if shared_mailbox:
            email_url = f"https://graph.microsoft.com/v1.0/users/{shared_mailbox}/messages/{message_id}"
        else:
            email_url = f"https://graph.microsoft.com/v1.0/me/messages/{message_id}"
        email_resp = requests.get(email_url, headers=headers)
        
        # If 401, try refreshing the token
        if email_resp.status_code == 401:
            print("[EstimateAnalysis] Token expired, attempting refresh...")
            refreshed = outlook_refresh_token(user_email)
            if refreshed:
                token = get_outlook_token()
                headers = {"Authorization": f"Bearer {token}"}
                email_resp = requests.get(email_url, headers=headers)
        
        if email_resp.status_code != 200:
            return jsonify({'error': f'Failed to fetch email: {email_resp.status_code}'}), 500
        
        email_data = email_resp.json()
        subject = email_data.get('subject', '')
        body = email_data.get('body', {}).get('content', '')
        from_name = email_data.get('from', {}).get('emailAddress', {}).get('name', '')
        from_email = email_data.get('from', {}).get('emailAddress', {}).get('address', '')
        
        # Clean HTML from body
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(body, 'html.parser')
        body_text = soup.get_text(separator='\n', strip=True)
        
        # Get attachments - use shared mailbox if applicable
        if shared_mailbox:
            attachments_url = f"https://graph.microsoft.com/v1.0/users/{shared_mailbox}/messages/{message_id}/attachments"
        else:
            attachments_url = f"https://graph.microsoft.com/v1.0/me/messages/{message_id}/attachments"
        att_resp = requests.get(attachments_url, headers=headers)
        attachments = att_resp.json().get('value', []) if att_resp.status_code == 200 else []
        
        # Extract content from PDF attachments - use Vision for images, PyPDF2 for text
        attachment_texts = []
        pdf_images = []  # For Claude Vision
        pdf_raw_bytes_list = []  # For hybrid text extraction
        
        for att in attachments:
            att_name = att.get('name', 'Unknown')
            att_type = att.get('contentType', '').lower()
            print(f"[EstimateAnalysis] Processing attachment: {att_name} ({att_type})")
            
            if att_type == 'application/pdf':
                try:
                    content_bytes_raw = base64.b64decode(att.get('contentBytes', ''))
                    pdf_raw_bytes_list.append(content_bytes_raw)  # Save for hybrid extraction
                    import io
                    import PyPDF2
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(content_bytes_raw))
                    pdf_text = ""
                    for page in pdf_reader.pages[:5]:  # First 5 pages max
                        pdf_text += page.extract_text() + "\n"
                    
                    print(f"[EstimateAnalysis] PDF text extracted: {len(pdf_text.strip())} chars")
                    
                    # Always use Vision for PDFs - it reads architectural plans better
                    print(f"[EstimateAnalysis] Converting PDF to images for Vision: {att_name}")
                    try:
                        import fitz  # PyMuPDF
                        print(f"[EstimateAnalysis] PyMuPDF imported successfully")
                        pdf_doc = fitz.open(stream=content_bytes_raw, filetype="pdf")
                        print(f"[EstimateAnalysis] PDF opened, has {len(pdf_doc)} pages")
                        for page_num in range(min(5, len(pdf_doc))):  # First 5 pages to capture floor plans with sq footage
                            page = pdf_doc[page_num]
                            # Render at 200 DPI for better clarity
                            mat = fitz.Matrix(200/72, 200/72)
                            pix = page.get_pixmap(matrix=mat)
                            
                            # Check if image exceeds 8000px and scale down if needed
                            if pix.width > 7000 or pix.height > 7000:
                                scale = 7000 / max(pix.width, pix.height)
                                mat = fitz.Matrix(scale * 200/72, scale * 200/72)
                                pix = page.get_pixmap(matrix=mat)
                            img_bytes = pix.tobytes("png")
                            img_base64 = base64.b64encode(img_bytes).decode('utf-8')
                            pdf_images.append({
                                'name': f"{att_name} - Page {page_num + 1}",
                                'data': img_base64
                            })
                            print(f"[EstimateAnalysis] Converted page {page_num + 1}")
                        pdf_doc.close()
                    except ImportError as ie:
                        print(f"[EstimateAnalysis] PyMuPDF not installed: {ie}")
                        # Fall back to text if Vision not available
                        if pdf_text.strip():
                            attachment_texts.append(f"=== PDF: {att_name} ===\n{pdf_text[:5000]}")
                    except Exception as img_err:
                        print(f"[EstimateAnalysis] Image extraction failed: {img_err}")
                        if pdf_text.strip():
                            attachment_texts.append(f"=== PDF: {att_name} ===\n{pdf_text[:5000]}")
                except Exception as e:
                    print(f"[EstimateAnalysis] PDF extraction failed: {e}")
        
        # Use Gemini + OCR analysis
        print(f"[EstimateAnalysis] PDF images: {len(pdf_images)}, Text attachments: {len(attachment_texts)}")

        if not gemini_estimate.is_ready():
            return jsonify({'error': 'Gemini AI or OCR not available'}), 500

        # Run Gemini + OCR analysis
        result = gemini_estimate.analyze_estimate(
            pdf_images=pdf_images,
            pdf_raw_bytes=pdf_raw_bytes_list,
            email_body=body_text,
            email_from=f"{from_name} <{from_email}>",
            email_subject=subject
        )
        
        if not result.get('success'):
            return jsonify({'error': result.get('error', 'Analysis failed')}), 500
        
        # Get the analysis from the result
        analysis = result.get('analysis', {})
        
        # Add email metadata
        analysis['email_from'] = from_name
        analysis['email_from_address'] = from_email
        analysis['email_subject'] = subject
        analysis['attachment_count'] = len(attachments)
        analysis['attachment_names'] = [a.get('name', '') for a in attachments]
        
        # Include pass details for debugging
        analysis['_passes'] = result.get('passes', {})
        
        print(f"[EstimateAnalysis] Complete: sqft={analysis.get('square_footage')}, type={analysis.get('project_type')}")
        
        return jsonify(analysis)
        
    except Exception as e:
        print(f"[EstimateAnalysis] Error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

# ==================== END ESTIMATE ANALYSIS ====================


if __name__ == '__main__':
    print("Starting ProCalcs API Server...")
    print("Dashboard will connect to: http://localhost:5000")
    # Auto-start email scanner
    email_scanner.start_scanner()
    app.run(debug=True, port=5000)
